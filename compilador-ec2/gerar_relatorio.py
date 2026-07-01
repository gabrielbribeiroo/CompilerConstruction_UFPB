"""Gera o arquivo RELATORIO.docx a partir do conteudo do RELATORIO.md.

Uso:
    python gerar_relatorio.py

Requer: python-docx (pip install python-docx).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT = "Calibri"
BODY_SIZE = Pt(11)
CODE_FONT = "Consolas"
HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x5F)
TABLE_HEADER_BG = "1F3A5F"
CODE_BG = "F4F4F4"


# -------- helpers --------

def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_paragraph_shading(paragraph, hex_color):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.bold = True
    if level == 0:
        run.font.size = Pt(20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = HEADING_COLOR
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = HEADING_COLOR
    else:
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(14 if level > 0 else 0)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_runs_with_inline_code(p, text):
    parts = text.split("`")
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.font.name = FONT
        run.font.size = BODY_SIZE
        if i % 2 == 1:
            run.font.name = CODE_FONT
            run.font.size = Pt(10)


def add_body(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(6)
    add_runs_with_inline_code(p, text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_runs_with_inline_code(p, text)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    set_paragraph_shading(p, CODE_BG)
    linhas = code.rstrip("\n").split("\n")
    for i, linha in enumerate(linhas):
        if i > 0:
            run = p.add_run()
            run.add_break()
        run = p.add_run(linha)
        run.font.name = CODE_FONT
        run.font.size = Pt(9.5)
    return p


def add_centered_info_line(doc, label, value, bold_label=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    if label:
        r = p.add_run(label + " ")
        r.bold = bold_label
        r.font.name = FONT
        r.font.size = BODY_SIZE
    r2 = p.add_run(value)
    r2.font.name = FONT
    r2.font.size = BODY_SIZE


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for idx, h in enumerate(headers):
        hdr_cells[idx].text = ""
        r = hdr_cells[idx].paragraphs[0].add_run(h)
        r.bold = True
        r.font.name = FONT
        r.font.size = BODY_SIZE
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[idx], TABLE_HEADER_BG)
        hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if col_widths:
            hdr_cells[idx].width = col_widths[idx]
    for row in rows:
        cells = table.add_row().cells
        for idx, valor in enumerate(row):
            cells[idx].text = ""
            add_runs_with_inline_code(cells[idx].paragraphs[0], str(valor))
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if col_widths:
                cells[idx].width = col_widths[idx]
    return table


# -------- documento --------

def build_document(path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = BODY_SIZE
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # Cabecalho institucional
    add_centered_info_line(doc, "", "Universidade Federal da Paraíba (UFPB)")
    add_centered_info_line(doc, "", "Centro de Informática – Curso de Ciência da Computação")
    add_centered_info_line(doc, "Disciplina:", "Construção de Compiladores 1")
    add_centered_info_line(doc, "Professor:", "Andrei de Araújo Formiga")
    doc.add_paragraph()

    # Titulo
    add_heading(doc, "Relatório – Atividade 07", 0)
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Compilador EC2 — Precedência e Associatividade")
    r_sub.italic = True
    r_sub.font.size = Pt(13)
    r_sub.font.name = FONT
    p_sub.paragraph_format.space_after = Pt(18)

    # Integrantes
    add_heading(doc, "Integrantes do grupo", 1)
    add_table(
        doc,
        headers=["Nome", "Matrícula"],
        rows=[
            ["Davi Alves Rodrigues", "20230102377"],
            ["Gabriel Barbosa Ribeiro de Oliveira", "20230012814"],
            ["João Vitor Sampaio Costa", "20230089776"],
            ["Nathan Meira Nóbrega", "20240008904"],
        ],
    )
    doc.add_paragraph()

    # =============== O QUE FOI IMPLEMENTADO ===============
    add_heading(doc, "O que foi implementado", 1)

    # 1. O problema
    add_heading(doc, "1. O problema: parênteses obrigatórios vs. notação natural", 2)
    add_body(
        doc,
        "Em EC1, toda operação binária precisa de parênteses explícitos "
        "(`(a op b)`), o que elimina qualquer ambiguidade mas é pouco "
        "natural para quem está acostumado com a notação matemática usual, "
        "onde `7 + 5 * 3` significa `7 + (5 * 3)` por convenção de "
        "precedência. O objetivo desta atividade é fazer o analisador "
        "sintático aceitar expressões nessa notação mais natural, sem "
        "exigir parênteses em toda operação.",
    )
    add_body(
        doc,
        "Simplesmente remover os parênteses da gramática de EC1 "
        "(`<exp> ::= <exp> <op> <exp>`) não funciona por dois motivos "
        "discutidos no enunciado:",
    )
    add_bullet(
        doc,
        "Ambiguidade de precedência — sem distinguir níveis, `7 + 5 * 3` "
        "teria duas árvores de derivação possíveis, uma calculando a soma "
        "primeiro e outra a multiplicação primeiro.",
    )
    add_bullet(
        doc,
        "Recursão à esquerda — a produção `<exp> ::= <exp> <op> <exp>`, "
        "mesmo com níveis de precedência separados, gera uma função de "
        "parser que chama a si mesma como primeira ação "
        "(`exp(): e1 = exp(); ...`), resultando em loop infinito em um "
        "parser descendente recursivo.",
    )

    # 2. Gramatica EC2
    add_heading(doc, "2. Gramática EC2", 2)
    add_body(
        doc,
        "A gramática usada, exatamente a do enunciado (seção 3), resolve "
        "os dois problemas simultaneamente: introduz um não-terminal por "
        "nível de precedência (resolve a ambiguidade) e usa a forma "
        "iterativa `(...)* ` em vez de recursão à esquerda direta (resolve "
        "o loop infinito e mantém associatividade à esquerda — ao "
        "contrário de simplesmente inverter a ordem para "
        "`<termo> '+' <exp>`, que geraria associatividade à direita e "
        "quebraria a semântica de `-` e `/`):",
    )
    add_code_block(
        doc,
        "<exp_a> ::= <exp_m> (('+' | '-') <exp_m>)*\n"
        "<exp_m> ::= <prim>  (('*' | '/') <prim>)*\n"
        "<prim>  ::= <num> | '(' <exp_a> ')'",
    )
    add_body(
        doc,
        "`exp_a` é o nível aditivo (precedência mais baixa), `exp_m` o "
        "multiplicativo (precedência mais alta), e `prim` reconhece uma "
        "constante ou reinicia a análise em `exp_a` dentro de parênteses.",
    )

    # 3. Parser
    add_heading(doc, "3. Parser (parser.py) — reescrito", 2)
    add_body(
        doc,
        "Implementamos uma função por não-terminal, seguindo o "
        "pseudocódigo do enunciado (seção 3) ao pé da letra:",
    )
    add_bullet(
        doc,
        "`_analisa_exp_a()` — reconhece um `exp_m` (operando esquerdo com "
        "`_analisa_exp_m()`), depois entra num laço: enquanto o próximo "
        "token (via `olhar_proximo()`, sem consumir) for `+` ou `-`, "
        "consome o operador, reconhece outro `exp_m`, e monta "
        "`esq = OpBin(op, esq, dir)` — o nó recém-criado vira o novo "
        "operando esquerdo da próxima iteração, o que produz "
        "associatividade à esquerda automaticamente, sem "
        "pós-processamento na árvore.",
    )
    add_bullet(
        doc,
        "`_analisa_exp_m()` — estrutura idêntica a `_analisa_exp_a`, mas "
        "para `*`/`/`, chamando `_analisa_prim()` em vez de "
        "`_analisa_exp_m()`.",
    )
    add_bullet(
        doc,
        "`_analisa_prim()` — se o próximo token for `NUMERO`, devolve "
        "`Const`; se for `PAREN_ESQ`, consome, chama `_analisa_exp_a()` "
        "recursivamente (reiniciando no nível de precedência mais baixo "
        "dentro dos parênteses) e exige `PAREN_DIR` ao final.",
    )
    add_bullet(
        doc,
        "`analisa_programa()` — chama `_analisa_exp_a()` e exige `EOF` em "
        "seguida, mesmo invariante usado desde a Atividade 05 para "
        "capturar lixo após a expressão raiz.",
    )
    add_body(
        doc,
        "O ponto de atenção citado no enunciado — “não retirar o token do "
        "fluxo de entrada ao verificar se é operador” — já estava "
        "resolvido, porque o `lexer.py` reusado (sem alteração) já expõe "
        "`olhar_proximo()` desde a Atividade 04/05 especificamente para "
        "esse tipo de decisão por lookahead sem consumo.",
    )

    # 4. Reuso
    add_heading(doc, "4. Reuso sem alteração de lexer, AST e codegen", 2)
    add_body(
        doc,
        "Confirmado com `diff` byte a byte: `lexer.py`, `ast_ec1.py`, "
        "`codegen.py` e `runtime.s` neste diretório são idênticos aos da "
        "Atividade 06. Isso é intencional e reflete diretamente o que o "
        "enunciado pede na seção 4 (“Artefato para entrega”): “A geração "
        "de código não muda com relação à Atividade 06.” A árvore de "
        "sintaxe abstrata (`Exp`, `Const`, `OpBin`, `Op`) também não "
        "precisa mudar — só muda como ela é construída, não sua "
        "representação. O analisador léxico já reconhecia todos os tokens "
        "necessários (`NUMERO`, `PAREN_ESQ`, `PAREN_DIR`, `SOMA`, `SUB`, "
        "`MULT`, `DIV`) desde a Atividade 04.",
    )

    # 5. CLI
    add_heading(doc, "5. Ponto de entrada (compec2.py)", 2)
    add_body(
        doc,
        "Estrutura idêntica ao `compec1.py` da Atividade 06 — apenas "
        "troca o import de `parser` (agora o de EC2) e a extensão "
        "reconhecida (`.ec2` → `.s`). Erros léxicos, sintáticos ou de E/S "
        "vão para `stderr` com exit code 1, sem gerar `.s`.",
    )

    # 6. Testes
    add_heading(doc, "6. Suíte de testes (tests/test_parser_precedencia.py)", 2)
    add_body(doc, "26 testes, 0 falhas, em 6 classes:")
    add_table(
        doc,
        headers=["Classe", "Foco", "Casos"],
        rows=[
            [
                "TestPrecedencia",
                "* e / calculados antes de + e -; parênteses sobrepõem a "
                "precedência padrão",
                "6",
            ],
            [
                "TestAssociatividade",
                "Associatividade à esquerda em +, -, *, /, e em cadeias longas",
                "5",
            ],
            [
                "TestCompatibilidadeComEC1",
                "Programas EC1 (parênteses obrigatórios) continuam aceitos "
                "e avaliando igual",
                "5",
            ],
            [
                "TestErrosSintaticos",
                "Detecção de entrada mal formada",
                "7",
            ],
            [
                "TestEquivalenciaSemantica",
                "Código gerado (codegen.py inalterado) bate com avaliar() "
                "para 13 programas",
                "1",
            ],
            [
                "TestCLI",
                "Comportamento de compec2.py como subprocesso",
                "2",
            ],
            ["Total", "", "26"],
        ],
    )
    doc.add_paragraph()
    add_body(doc, "Destaques:")
    add_bullet(
        doc,
        "`test_soma_e_multiplicacao_estrutura` verifica que a AST de "
        "`7 + 5 * 3` é exatamente "
        "`OpBin(SOMA, Const(7), OpBin(MULT, Const(5), Const(3)))` — não a "
        "árvore alternativa que calcularia a soma primeiro.",
    )
    add_bullet(
        doc,
        "`test_subtracao_associa_a_esquerda_estrutura` e "
        "`test_divisao_associa_a_esquerda_estrutura` verificam a árvore "
        "literal de `10 - 8 - 2` e `100 / 10 / 2`, garantindo que a "
        "associatividade está correta na estrutura, não só coincidindo "
        "no valor final por acaso.",
    )
    add_bullet(
        doc,
        "`TestEquivalenciaSemantica` reaproveita a técnica da Atividade "
        "06 — um simulador da máquina de pilha em Python que interpreta "
        "as instruções emitidas pelo `codegen.py` — para confirmar que o "
        "mesmo gerador de código, sem nenhuma alteração, continua "
        "correto quando alimentado com árvores produzidas pelo novo "
        "parser de precedência.",
    )

    # 7. Exemplos
    add_heading(doc, "7. Arquivos de exemplo", 2)
    add_body(
        doc,
        "Seis programas válidos e um inválido em `exemplos/`, cobrindo "
        "precedência (`7 + 5 * 3`), parênteses sobrepondo a precedência "
        "padrão (`(7 + 5) * 3`), associatividade em subtração e divisão "
        "(`10 - 8 - 2`, `100 / 10 / 2`), uma expressão mista com os "
        "quatro operadores (`2 + 3 * 4 - 10 / 5`), compatibilidade "
        "retroativa com um programa EC1 do enunciado da Atividade 04 "
        "(`(33 + (912 * 11))`), e um erro sintático proposital "
        "(`7 + * 3`, operador sem operando).",
    )
    add_table(
        doc,
        headers=["Arquivo", "Conteúdo", "Resultado"],
        rows=[
            ["valido1.ec2", "7 + 5 * 3", "22"],
            ["valido2.ec2", "(7 + 5) * 3", "36"],
            ["valido3.ec2", "10 - 8 - 2", "0"],
            ["valido4.ec2", "100 / 10 / 2", "5"],
            ["valido5.ec2", "2 + 3 * 4 - 10 / 5", "12"],
            ["valido6.ec2", "(33 + (912 * 11))", "10065"],
            ["invalido1.ec2", "7 + * 3", "erro sintático, exit 1"],
        ],
    )

    # =============== ESTRUTURA DE ARQUIVOS ===============
    add_heading(doc, "Estrutura de arquivos entregue", 1)
    add_code_block(
        doc,
        "compilador-ec2/\n"
        "├── lexer.py           # idêntico à Atividade 06\n"
        "├── ast_ec1.py         # idêntico à Atividade 06\n"
        "├── parser.py          # NOVO\n"
        "├── codegen.py         # idêntico à Atividade 06\n"
        "├── compec2.py\n"
        "├── runtime.s          # idêntico à Atividade 06\n"
        "├── exemplos/\n"
        "│   ├── valido1.ec2\n"
        "│   ├── valido2.ec2\n"
        "│   ├── valido3.ec2\n"
        "│   ├── valido4.ec2\n"
        "│   ├── valido5.ec2\n"
        "│   ├── valido6.ec2\n"
        "│   └── invalido1.ec2\n"
        "├── tests/\n"
        "│   └── test_parser_precedencia.py\n"
        "├── README.md\n"
        "├── PLANO.md\n"
        "└── RELATORIO.md",
    )

    # =============== DECISOES DE PROJETO ===============
    add_heading(doc, "Decisões de projeto", 1)

    add_heading(
        doc,
        "Por que a forma (...)* em vez de recursão à direita "
        "(<termo> '+' <exp>)?",
        2,
    )
    add_body(
        doc,
        "A recursão à direita resolveria o loop infinito, mas geraria "
        "associatividade à direita, que é semanticamente errada para `-` "
        "e `/`: `10 - 8 - 2` viraria `10 - (8 - 2) = 4` em vez do correto "
        "`0`. A forma iterativa do enunciado, onde o nó recém-criado vira "
        "o novo operando esquerdo a cada iteração do laço, dá "
        "associatividade à esquerda diretamente, sem exigir nenhum "
        "pós-processamento ou inversão da árvore.",
    )

    add_heading(
        doc,
        "Por que usar olhar_proximo() (peek) em vez de consumir e "
        "“devolver” o token?",
        2,
    )
    add_body(
        doc,
        "O `AnalisadorLexico` reusado já expõe exatamente essa interface "
        "desde a Atividade 04/05, então não havia motivo para reinventar. "
        "Consumir e tentar “devolver” um token exigiria um mecanismo de "
        "pushback adicional no lexer; o peek evita esse problema porque "
        "nunca avança o cursor até confirmar que o token pertence à "
        "produção atual.",
    )

    add_heading(
        doc,
        "Por que reusar codegen.py sem qualquer edição, incluindo o "
        "comentário que ainda menciona “EC1”?",
        2,
    )
    add_body(
        doc,
        "O enunciado afirma explicitamente que a geração de código não "
        "muda em relação à Atividade 06. Reusar o arquivo byte a byte "
        "idêntico (e verificado com `diff`) é a demonstração mais direta "
        "de que essa afirmação é verdadeira: o compilador EC2 herda a "
        "mesma AST e o mesmo esquema de tradução baseado em pilha, e a "
        "única mudança de fato está concentrada em `parser.py`. Editar o "
        "comentário só para dizer “EC2” seria puramente cosmético e "
        "romperia essa garantia de identidade sem nenhum ganho — por "
        "isso optamos por documentar a decisão aqui em vez de tocar no "
        "arquivo.",
    )

    add_heading(
        doc,
        "Por que os testes de estrutura da AST (não só de valor) são "
        "importantes?",
        2,
    )
    add_body(
        doc,
        "Um teste que só verifica o valor final pode passar por "
        "acidente — por exemplo, `7 + 5 * 3` daria `36` só se a soma "
        "fosse calculada primeiro E o resultado fosse coincidentemente "
        "igual ao valor esperado em outro caso. Verificar a árvore "
        "literalmente (`OpBin(SOMA, Const(7), OpBin(MULT, Const(5), "
        "Const(3)))`) garante que a estrutura da precedência e "
        "associatividade está correta, não só que o número final bateu "
        "por coincidência aritmética.",
    )

    # =============== DIFICULDADES ===============
    add_heading(doc, "Dificuldades", 1)
    add_body(
        doc,
        "Nenhuma dificuldade significativa. O pseudocódigo do enunciado "
        "(seção 3) já descreve o parser praticamente pronto para "
        "tradução direta em Python; o único cuidado foi generalizar a "
        "mesma estrutura de laço para `exp_m` sem duplicar lógica "
        "desnecessária (as duas funções, `_analisa_exp_a` e "
        "`_analisa_exp_m`, têm exatamente a mesma forma, diferindo apenas "
        "no dicionário de operadores e na função de nível inferior que "
        "chamam — conforme o próprio enunciado observa na seção 3).",
    )

    doc.save(path)


if __name__ == "__main__":
    import os
    out = os.path.join(os.path.dirname(__file__), "RELATORIO.docx")
    build_document(out)
    print(f"gerado: {out}")
