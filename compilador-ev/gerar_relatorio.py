"""Gera o arquivo RELATORIO.docx a partir do conteudo do RELATORIO.md.

Uso:
    python gerar_relatorio.py

Requer: python-docx (pip install python-docx).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT = "Calibri"
BODY_SIZE = Pt(11)
CODE_FONT = "Consolas"
HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x5F)
TABLE_HEADER_BG = "1F3A5F"
CODE_BG = "F4F4F4"


# -------- helpers --------

def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_paragraph_shading(paragraph, hex_color):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.bold = True
    if level == 0:
        run.font.size = Pt(20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = HEADING_COLOR
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = HEADING_COLOR
    else:
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(14 if level > 0 else 0)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_runs_with_inline_code(p, text):
    parts = text.split("`")
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.font.name = FONT
        run.font.size = BODY_SIZE
        if i % 2 == 1:
            run.font.name = CODE_FONT
            run.font.size = Pt(10)


def add_body(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(6)
    add_runs_with_inline_code(p, text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_runs_with_inline_code(p, text)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    set_paragraph_shading(p, CODE_BG)
    linhas = code.rstrip("\n").split("\n")
    for i, linha in enumerate(linhas):
        if i > 0:
            run = p.add_run()
            run.add_break()
        run = p.add_run(linha)
        run.font.name = CODE_FONT
        run.font.size = Pt(9.5)
    return p


def add_centered_info_line(doc, label, value, bold_label=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    if label:
        r = p.add_run(label + " ")
        r.bold = bold_label
        r.font.name = FONT
        r.font.size = BODY_SIZE
    r2 = p.add_run(value)
    r2.font.name = FONT
    r2.font.size = BODY_SIZE


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for idx, h in enumerate(headers):
        hdr_cells[idx].text = ""
        r = hdr_cells[idx].paragraphs[0].add_run(h)
        r.bold = True
        r.font.name = FONT
        r.font.size = BODY_SIZE
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[idx], TABLE_HEADER_BG)
        hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if col_widths:
            hdr_cells[idx].width = col_widths[idx]
    for row in rows:
        cells = table.add_row().cells
        for idx, valor in enumerate(row):
            cells[idx].text = ""
            add_runs_with_inline_code(cells[idx].paragraphs[0], str(valor))
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if col_widths:
                cells[idx].width = col_widths[idx]
    return table


# -------- documento --------

def build_document(path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = BODY_SIZE
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # Cabecalho institucional
    add_centered_info_line(doc, "", "Universidade Federal da Paraíba (UFPB)")
    add_centered_info_line(doc, "", "Centro de Informática – Curso de Ciência da Computação")
    add_centered_info_line(doc, "Disciplina:", "Construção de Compiladores 1")
    add_centered_info_line(doc, "Professor:", "Andrei de Araújo Formiga")
    doc.add_paragraph()

    # Titulo
    add_heading(doc, "Relatório – Atividade 08", 0)
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Compilador EV — Expressões com Variáveis")
    r_sub.italic = True
    r_sub.font.size = Pt(13)
    r_sub.font.name = FONT
    p_sub.paragraph_format.space_after = Pt(18)

    # Integrantes
    add_heading(doc, "Integrantes do grupo", 1)
    add_table(
        doc,
        headers=["Nome", "Matrícula"],
        rows=[
            ["Davi Alves Rodrigues", "20230102377"],
            ["Gabriel Barbosa Ribeiro de Oliveira", "20230012814"],
            ["João Vitor Sampaio Costa", "20230089776"],
            ["Nathan Meira Nóbrega", "20240008904"],
        ],
    )
    doc.add_paragraph()

    # =============== O QUE FOI IMPLEMENTADO ===============
    add_heading(doc, "O que foi implementado", 1)

    # 1. A linguagem EV
    add_heading(doc, "1. A linguagem EV e o novo estágio de análise semântica", 2)
    add_body(
        doc,
        "EV (Expressões com Variáveis) estende EC2 (Atividade 07) "
        "permitindo declarar e usar variáveis. Um programa é "
        "`<decl>* <result>`: zero ou mais declarações `nome = exp;`, "
        "seguidas de uma expressão final obrigatória `= exp`. Diferente "
        "das atividades anteriores — onde qualquer erro de "
        "“significado” já não existia (expressões constantes não têm "
        "como estar semanticamente erradas além de sintaxe/léxico) — "
        "EV introduz um tipo de erro que não é sintático: usar uma "
        "variável que nunca foi declarada é uma sentença perfeitamente "
        "aceita pela gramática, mas sem sentido. Isso exige uma nova "
        "etapa depois do parser, chamada de análise semântica (ou "
        "contextual), conforme o diagrama do enunciado:",
    )
    add_code_block(
        doc,
        "código-fonte → Análise Léxica → tokens → Análise Sintática → AST\n"
        "             → Análise Semântica → AST (verificada) → Geração de Código → assembly",
    )
    add_body(
        doc,
        "Diferente da Atividade 07 (onde lexer, AST e codegen eram "
        "herdados sem qualquer alteração da atividade anterior), aqui "
        "todos os estágios do compilador precisaram mudar, exatamente "
        "como o enunciado antecipa na introdução.",
    )

    # 2. Alteracoes lexicas
    add_heading(doc, "2. Alterações léxicas (lexer.py)", 2)
    add_body(
        doc,
        "Três tokens novos: `IDENT` (identificador), `IGUAL` (`=`, "
        "usado tanto para atribuição quanto para marcar a expressão "
        "final) e `PONTO_VIRGULA` (`;`, fim de declaração). Um "
        "identificador começa com uma letra e é seguido de zero ou "
        "mais letras/dígitos.",
    )
    add_body(
        doc,
        "Um detalhe da seção 3 do enunciado exigiu atenção: uma "
        "sequência de dígitos seguida imediatamente por uma letra sem "
        "separador (o exemplo do enunciado é `237axy`) continua sendo "
        "um erro léxico, não dois tokens (`237` e `axy`). Implementamos "
        "isso em `_ler_numero()`: depois de consumir os dígitos, se o "
        "próximo caractere for uma letra, levantamos `ErroLexico` na "
        "posição dessa letra, em vez de deixar o scanner seguir e "
        "tokenizar `axy` como um identificador separado.",
    )

    # 3. Alteracoes na sintaxe
    add_heading(doc, "3. Alterações na sintaxe (parser.py)", 2)
    add_body(
        doc,
        "A parte de expressões (`exp`/`exp_m`) é exatamente a mesma "
        "técnica da Atividade 07 — um não-terminal por nível de "
        "precedência, laço para as produções repetidas, "
        "associatividade à esquerda automática. A única mudança em "
        "`prim` é reconhecer `IDENT` como uma terceira alternativa, "
        "devolvendo um nó `Var`.",
    )
    add_body(
        doc,
        "O topo da gramática é novo: `analisa_programa()` segue o "
        "pseudocódigo do enunciado (seção 4) — olha o próximo token "
        "(via `olhar_proximo()`, sem consumir) e, enquanto for `IDENT`, "
        "reconhece uma declaração completa (`_analisa_decl()`); ao "
        "encontrar `=`, reconhece a expressão final e monta o nó raiz "
        "`Programa(declaracoes, exp_final)`. `_analisa_decl()` "
        "reconhece `ident '=' exp ';'` na ordem, devolvendo um nó "
        "`Decl(nome, exp)`.",
    )

    # 4. AST
    add_heading(doc, "4. Árvore de sintaxe abstrata (ast_ev.py)", 2)
    add_body(
        doc,
        "Reaproveita `Exp`/`Const`/`OpBin`/`Op` na mesma forma das "
        "atividades anteriores, e adiciona:",
    )
    add_bullet(
        doc,
        "`Var(nome, posicao)` — folha que referencia uma variável. O "
        "campo `posicao` usa `field(compare=False)` do `dataclasses`, "
        "para guardar a posição no texto-fonte (útil nas mensagens de "
        "erro semântico) sem que ela interfira na igualdade estrutural "
        "usada nos testes — `Var(\"x\")` e `Var(\"x\", posicao=17)` são "
        "iguais.",
    )
    add_bullet(
        doc,
        "`Decl(nome, exp)` — não é uma `Exp`; representa uma "
        "declaração.",
    )
    add_bullet(
        doc,
        "`Programa(declaracoes, exp_final)` — nó raiz. Ganhou um "
        "método `avaliar()` (sem argumentos) que monta um ambiente "
        "(`dict`) processando as declarações em ordem e por fim avalia "
        "a expressão final — o interpretador do programa inteiro.",
    )
    add_body(
        doc,
        "O método `avaliar()` de `Exp` passou a receber um parâmetro "
        "`env` (dicionário nome→valor): `Const.avaliar(env)` ignora "
        "`env`, `Var.avaliar(env)` faz `env[self.nome]`, e "
        "`OpBin.avaliar(env)` repassa `env` recursivamente aos dois "
        "filhos.",
    )

    # 5. Semantica
    add_heading(doc, "5. Análise semântica (semantica.py)", 2)
    add_body(doc, "Módulo novo com duas peças:")
    add_bullet(
        doc,
        "`TabelaSimbolos` — envolve um `set[str]` dos nomes já "
        "declarados, com `declarar(nome)` e `esta_declarada(nome)`. "
        "Para EV só precisamos saber se a variável existe — não há "
        "tipos a verificar (o próprio enunciado observa que "
        "verificação de tipos não se aplica aqui, pois todas as "
        "variáveis têm o mesmo tipo).",
    )
    add_bullet(
        doc,
        "`verifica_programa(programa)` — percorre "
        "`programa.declaracoes` na ordem em que aparecem no "
        "código-fonte: para cada declaração, verifica recursivamente "
        "(`_verifica_exp`) que toda `Var` na expressão já está na "
        "tabela; só então declara a variável dessa `Decl`. Ao final, "
        "faz a mesma verificação na expressão final. Levanta "
        "`ErroSemantico(nome, posicao)` no primeiro uso de variável não "
        "declarada — o processo para imediatamente, conforme a seção 5 "
        "do enunciado (“o compilador deve reportar um erro e parar”).",
    )
    add_body(
        doc,
        "Isso captura corretamente os dois erros do exemplo do "
        "enunciado (`x = 7 + y; y = x * 11; = x * y + z`): a "
        "referência a `y` dentro da declaração de `x` é rejeitada "
        "porque `y` ainda não foi declarada naquele ponto (mesmo que "
        "`y` seja declarada logo depois) — a ordem de varredura é "
        "sequencial, não “olha o programa inteiro primeiro”.",
    )

    # 6. Codegen
    add_heading(doc, "6. Geração de código (codegen.py)", 2)
    add_body(
        doc,
        "Reaproveita o esquema de pilha das Atividades 06/07 para "
        "`Const`/`OpBin`, e adiciona:",
    )
    add_bullet(
        doc,
        "`Var(nome)` → `mov <nome>, %rax` (lê a variável da memória "
        "para o registrador de resultado).",
    )
    add_bullet(
        doc,
        "Coleta de símbolos (`_coleta_variaveis`) — varre "
        "`programa.declaracoes` e monta a lista de nomes únicos, na "
        "ordem da primeira declaração, usando um `dict` como conjunto "
        "ordenado para deduplicar sem alterar a ordem (uma variável "
        "reatribuída mais de uma vez só gera uma diretiva `.lcomm`, "
        "evitando símbolo duplicado no assembly gerado).",
    )
    add_bullet(
        doc,
        "Seção `.bss` — uma diretiva `.lcomm <nome>, 8` por variável "
        "(8 bytes = inteiro de 64 bits), inserida antes da seção "
        "`.text`, exatamente como no modelo da seção 6 do enunciado.",
    )
    add_bullet(
        doc,
        "Código de cada declaração — o código da expressão do lado "
        "direito (o mesmo esquema de pilha de sempre), seguido de "
        "`mov %rax, <nome>` para gravar o resultado na variável.",
    )
    add_bullet(
        doc,
        "Modelo de saída — agora tem duas partes preenchidas: a seção "
        "`.bss` (variáveis) e a seção `.text` (código de cada "
        "declaração, em ordem, seguido do código da expressão final).",
    )
    add_body(
        doc,
        "Também adicionamos comentários (`# nome = exp;`) antes do "
        "código de cada declaração e `# = exp` antes da expressão "
        "final, usando o `__str__()` das classes da AST — um recurso "
        "barato que deixa o `.s` gerado mais fácil de ler durante o "
        "desenvolvimento.",
    )

    # 7. CLI
    add_heading(doc, "7. Ponto de entrada (compev.py)", 2)
    add_body(
        doc,
        "Pipeline completo: lê o arquivo `.ev` → `analisar()` (léxico "
        "+ sintático) → `verifica_programa()` (semântico) → "
        "`gerar_programa()` (codegen) → grava `.s`. Qualquer uma das "
        "três exceções (`ErroLexico`, `ErroSintatico`, "
        "`ErroSemantico`) é capturada no mesmo bloco, imprime a "
        "mensagem em `stderr` e encerra com exit code 1, sem gravar "
        "`.s`.",
    )

    # 8. Variacao sintatica
    add_heading(doc, "8. Variação sintática", 2)
    add_body(
        doc,
        "Seguimos exatamente a sintaxe do enunciado (seção 2), sem "
        "adotar nenhuma das variações citadas na seção 7 (palavra-chave "
        "`return`, sintaxe estilo Pascal com `var`/`begin`/`end`, "
        "sintaxe estilo C com tipos e `main()`) — apresentadas "
        "explicitamente como opcionais, “de acordo com o interesse do "
        "grupo”.",
    )

    # 9. Testes
    add_heading(doc, "9. Suíte de testes (tests/test_ev.py)", 2)
    add_body(doc, "32 testes, 0 falhas, em 7 classes:")
    add_table(
        doc,
        headers=["Classe", "Foco", "Casos"],
        rows=[
            ["TestLexico", "Tokens novos e o erro léxico dígito+letra", "5"],
            ["TestParser", "Estrutura da AST (Programa/Decl/Var)", "4"],
            ["TestErrosSintaticos", "Entradas mal formadas", "5"],
            [
                "TestSemantica",
                "Os dois erros do exemplo do enunciado + casos válidos",
                "5",
            ],
            ["TestInterpretacao", "Programa.avaliar()", "4"],
            [
                "TestEquivalenciaSemantica",
                "Código gerado bate com o interpretador",
                "1",
            ],
            [
                "TestCodegen",
                "Seção .bss, dedup de .lcomm, modelo completo",
                "4",
            ],
            ["TestCLI", "Comportamento de compev.py como subprocesso", "4"],
            ["Total", "", "32"],
        ],
    )
    doc.add_paragraph()
    add_body(doc, "Destaques:")
    add_bullet(
        doc,
        "`test_uso_de_variavel_antes_da_propria_declaracao` e "
        "`test_variavel_nao_declarada_na_expressao_final` reproduzem "
        "literalmente os dois erros do exemplo da seção 5 do "
        "enunciado, e verificam que o nome da variável relatado no "
        "`ErroSemantico` é exatamente o esperado (`y` e `w`, "
        "respectivamente).",
    )
    add_bullet(
        doc,
        "`test_exemplo_completo_do_enunciado` avalia o programa de "
        "exemplo completo (`x`, `y` e a expressão final) e confirma o "
        "resultado `60467` citado no enunciado.",
    )
    add_bullet(
        doc,
        "`TestEquivalenciaSemantica` estende o simulador de máquina de "
        "pilha das Atividades 06/07 com um dicionário de memória "
        "(`mem`), simulando também as instruções `mov %rax, <var>` e "
        "`mov <var>, %rax`, e compara o resultado com "
        "`Programa.avaliar()` para 6 programas, incluindo reatribuição "
        "de variável e divisões encadeadas.",
    )
    add_bullet(
        doc,
        "`TestCodegen.test_variavel_reatribuida_nao_duplica_lcomm` "
        "confirma que reatribuir uma variável (`x = 10; x = x + 5;`) "
        "gera apenas uma diretiva `.lcomm x, 8`, não duas (o que seria "
        "um símbolo duplicado e provavelmente um erro de montagem).",
    )

    # 10. Exemplos
    add_heading(doc, "10. Arquivos de exemplo", 2)
    add_body(
        doc,
        "Quatro programas válidos — o exemplo do perímetro do "
        "enunciado (`l + l + c + c`, resultado `140`), o exemplo "
        "completo do enunciado (resultado `60467`), um programa sem "
        "nenhuma declaração (só a expressão final) e um programa com "
        "reatribuição da mesma variável — e três inválidos: os dois "
        "erros semânticos do exemplo do enunciado (separados em dois "
        "arquivos, cada um isolando um dos dois problemas) e um erro "
        "léxico (`237axy`).",
    )
    add_table(
        doc,
        headers=["Arquivo", "Conteúdo", "Resultado"],
        rows=[
            ["valido1.ev", "l=30; c=40; = l+l+c+c", "140"],
            [
                "valido2.ev",
                "x=(7+4)*12; y=x*3+11; = (x*y)+(x*11)+(y*13)",
                "60467",
            ],
            ["valido3.ev", "= 6 * 7 (sem declarações)", "42"],
            ["valido4.ev", "x=10; x=x+5; = x (reatribuição)", "15"],
            [
                "invalido_uso_antes_da_declaracao.ev",
                "x=7+y; y=x*11; = x*y+z",
                "erro semântico (y), exit 1",
            ],
            [
                "invalido_variavel_nao_declarada.ev",
                "x=7*8; = x+w",
                "erro semântico (w), exit 1",
            ],
            [
                "invalido_erro_lexico_num_letra.ev",
                "= 237axy",
                "erro léxico, exit 1",
            ],
        ],
    )

    # =============== EXEMPLO DE SAIDA ===============
    add_heading(doc, "Exemplo de saída — perímetro do retângulo", 1)
    add_code_block(
        doc,
        "#\n"
        "# Saida gerada pelo compilador EV\n"
        "#\n"
        "    .section .bss\n"
        "    .lcomm l, 8\n"
        "    .lcomm c, 8\n"
        "    .section .text\n"
        "    .globl _start\n"
        "\n"
        "_start:\n"
        "    # l = 30;\n"
        "    mov $30, %rax\n"
        "    mov %rax, l\n"
        "    # c = 40;\n"
        "    mov $40, %rax\n"
        "    mov %rax, c\n"
        "    # = (((l + l) + c) + c)\n"
        "    mov c, %rax\n"
        "    push %rax\n"
        "    mov c, %rax\n"
        "    push %rax\n"
        "    mov l, %rax\n"
        "    push %rax\n"
        "    mov l, %rax\n"
        "    pop %rbx\n"
        "    add %rbx, %rax\n"
        "    pop %rbx\n"
        "    add %rbx, %rax\n"
        "    pop %rbx\n"
        "    add %rbx, %rax\n"
        "    call imprime_num\n"
        "    call sair\n"
        "\n"
        '    .include "runtime.s"',
    )
    add_body(doc, "Executado, imprime `140`.")

    # =============== ESTRUTURA DE ARQUIVOS ===============
    add_heading(doc, "Estrutura de arquivos entregue", 1)
    add_code_block(
        doc,
        "compilador-ev/\n"
        "├── lexer.py\n"
        "├── ast_ev.py\n"
        "├── parser.py\n"
        "├── semantica.py\n"
        "├── codegen.py\n"
        "├── compev.py\n"
        "├── runtime.s\n"
        "├── exemplos/\n"
        "│   ├── valido1.ev\n"
        "│   ├── valido2.ev\n"
        "│   ├── valido3.ev\n"
        "│   ├── valido4.ev\n"
        "│   ├── invalido_uso_antes_da_declaracao.ev\n"
        "│   ├── invalido_variavel_nao_declarada.ev\n"
        "│   └── invalido_erro_lexico_num_letra.ev\n"
        "├── tests/\n"
        "│   └── test_ev.py\n"
        "├── README.md\n"
        "├── PLANO.md\n"
        "└── RELATORIO.md",
    )

    # =============== DECISOES DE PROJETO ===============
    add_heading(doc, "Decisões de projeto", 1)

    add_heading(doc, "Por que Var.posicao usa field(compare=False)?", 2)
    add_body(
        doc,
        "Precisávamos guardar a posição do identificador para produzir "
        "mensagens de erro semântico úteis (`Erro semantico na "
        "posicao N: ...`), mas os testes de estrutura da AST comparam "
        "nós construídos manualmente (`Var(\"x\")`) com os produzidos "
        "pelo parser (`Var(\"x\", posicao=17)`). Marcar o campo como "
        "`compare=False` resolve os dois lados: a posição fica "
        "disponível quando existe, e a igualdade estrutural (usada nos "
        "testes e potencialmente em otimizações futuras) ignora um "
        "detalhe que não afeta o significado do nó.",
    )

    add_heading(
        doc,
        "Por que a análise semântica processa as declarações em ordem "
        "sequencial, e não “coleta todos os nomes primeiro”?",
        2,
    )
    add_body(
        doc,
        "Porque é isso que o enunciado especifica e é isso que produz "
        "o comportamento correto do exemplo dado: `x = 7 + y;` deve "
        "ser um erro, mesmo que `y` seja declarada na linha seguinte. "
        "Se coletássemos todos os nomes de uma vez antes de verificar "
        "qualquer expressão, esse programa passaria (incorretamente) "
        "na verificação semântica. A diferença entre as duas "
        "abordagens é explicitamente discutida no enunciado (o "
        "parêntese que menciona “em algumas linguagens” a declaração "
        "de `x` seria permitida) — escolhemos a semântica de EV (uso "
        "só após declaração), não a alternativa mencionada.",
    )

    add_heading(
        doc,
        "Por que deduplicar variáveis na seção .bss em vez de permitir "
        ".lcomm duplicado?",
        2,
    )
    add_body(
        doc,
        "O enunciado não proíbe explicitamente reatribuir uma variável "
        "já declarada (`x = 10; x = x + 5;` é sintaticamente válido e "
        "semanticamente correto — `x` já está declarada quando a "
        "segunda linha usa `x` na expressão). Sem dedup, o codegen "
        "emitiria duas diretivas `.lcomm x, 8` para o mesmo símbolo, o "
        "que é um símbolo duplicado e pode falhar na montagem "
        "dependendo do assembler. Deduplicar (mantendo a ordem da "
        "primeira declaração) resolve isso sem custo adicional.",
    )

    add_heading(doc, "Por que os comentários # nome = exp; no código gerado?", 2)
    add_body(
        doc,
        "Custo desprezível (usa o `__str__()` que a AST já precisa ter "
        "para depuração/testes) e torna o `.s` gerado muito mais fácil "
        "de inspecionar manualmente durante o desenvolvimento — foi "
        "assim que detectamos rapidamente, olhando a saída, que a "
        "ordem da pilha estava correta antes mesmo de escrever o teste "
        "de equivalência semântica.",
    )

    add_heading(
        doc,
        "Por que reusar o esquema de pilha e o mapeamento de "
        "operadores exatamente como nas Atividades 06/07, em vez de "
        "adotar mul de operando único como no exemplo do enunciado?",
        2,
    )
    add_body(
        doc,
        "O enunciado usa `mul %rbx` no seu exemplo ilustrativo de "
        "geração de código (seção 6), mas isso é apenas um exemplo — o "
        "“Artefato para entrega” (seção 8) pede apenas que a geração "
        "seja correta, sem exigir uma instrução específica. Mantivemos "
        "`imul %rbx, %rax` (forma de dois operandos) pela mesma razão "
        "documentada nas atividades anteriores: não sobrescreve "
        "`%rdx` desnecessariamente, e mantém o compilador consistente "
        "com a divisão com sinal (`cqo` + `idiv`), que já usa `%rdx` "
        "para outra finalidade.",
    )

    # =============== DIFICULDADES ===============
    add_heading(doc, "Dificuldades", 1)
    add_body(
        doc,
        "Nenhuma dificuldade significativa. A estrutura geral (lexer "
        "→ parser → nova etapa semântica → codegen) é bem guiada pelo "
        "enunciado. Os pontos que exigiram mais atenção:",
    )
    add_bullet(
        doc,
        "Erro léxico de dígito seguido de letra. Não é o "
        "comportamento “óbvio” de um scanner ingênuo (que tokenizaria "
        "`237` e `axy` separadamente); precisou de uma checagem "
        "explícita logo após consumir os dígitos do número.",
    )
    add_bullet(
        doc,
        "Ordem de verificação semântica. Garantir que a verificação "
        "acontece declaração por declaração, na ordem do "
        "arquivo-fonte — e não “verificar tudo de uma vez com todos os "
        "nomes disponíveis” — para reproduzir corretamente o "
        "comportamento esperado no exemplo do enunciado.",
    )
    add_bullet(
        doc,
        "Deduplicação de símbolos na seção .bss para o caso (não "
        "proibido pela gramática) de uma variável ser reatribuída mais "
        "de uma vez no mesmo programa.",
    )

    doc.save(path)


if __name__ == "__main__":
    import os
    out = os.path.join(os.path.dirname(__file__), "RELATORIO.docx")
    build_document(out)
    print(f"gerado: {out}")
