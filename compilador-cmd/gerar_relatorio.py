"""Gera o arquivo RELATORIO.docx a partir do conteudo do RELATORIO.md.

Uso:
    python gerar_relatorio.py

Requer: python-docx (pip install python-docx).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT = "Calibri"
BODY_SIZE = Pt(11)
CODE_FONT = "Consolas"
HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x5F)
TABLE_HEADER_BG = "1F3A5F"
CODE_BG = "F4F4F4"


# -------- helpers --------

def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_paragraph_shading(paragraph, hex_color):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.bold = True
    if level == 0:
        run.font.size = Pt(20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = HEADING_COLOR
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = HEADING_COLOR
    else:
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(14 if level > 0 else 0)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_runs_with_inline_code(p, text):
    parts = text.split("`")
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.font.name = FONT
        run.font.size = BODY_SIZE
        if i % 2 == 1:
            run.font.name = CODE_FONT
            run.font.size = Pt(10)


def add_body(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(6)
    add_runs_with_inline_code(p, text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_runs_with_inline_code(p, text)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    set_paragraph_shading(p, CODE_BG)
    linhas = code.rstrip("\n").split("\n")
    for i, linha in enumerate(linhas):
        if i > 0:
            run = p.add_run()
            run.add_break()
        run = p.add_run(linha)
        run.font.name = CODE_FONT
        run.font.size = Pt(9.5)
    return p


def add_centered_info_line(doc, label, value, bold_label=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    if label:
        r = p.add_run(label + " ")
        r.bold = bold_label
        r.font.name = FONT
        r.font.size = BODY_SIZE
    r2 = p.add_run(value)
    r2.font.name = FONT
    r2.font.size = BODY_SIZE


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for idx, h in enumerate(headers):
        hdr_cells[idx].text = ""
        r = hdr_cells[idx].paragraphs[0].add_run(h)
        r.bold = True
        r.font.name = FONT
        r.font.size = BODY_SIZE
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[idx], TABLE_HEADER_BG)
        hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if col_widths:
            hdr_cells[idx].width = col_widths[idx]
    for row in rows:
        cells = table.add_row().cells
        for idx, valor in enumerate(row):
            cells[idx].text = ""
            add_runs_with_inline_code(cells[idx].paragraphs[0], str(valor))
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if col_widths:
                cells[idx].width = col_widths[idx]
    return table


# -------- documento --------

def build_document(path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = BODY_SIZE
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # Cabecalho institucional
    add_centered_info_line(doc, "", "Universidade Federal da Paraíba (UFPB)")
    add_centered_info_line(doc, "", "Centro de Informática – Curso de Ciência da Computação")
    add_centered_info_line(doc, "Disciplina:", "Construção de Compiladores 1")
    add_centered_info_line(doc, "Professor:", "Andrei de Araújo Formiga")
    doc.add_paragraph()

    # Titulo
    add_heading(doc, "Relatório – Atividade 09", 0)
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Compilador Cmd — Comandos")
    r_sub.italic = True
    r_sub.font.size = Pt(13)
    r_sub.font.name = FONT
    p_sub.paragraph_format.space_after = Pt(18)

    # Integrantes
    add_heading(doc, "Integrantes do grupo", 1)
    add_table(
        doc,
        headers=["Nome", "Matrícula"],
        rows=[
            ["Davi Alves Rodrigues", "20230102377"],
            ["Gabriel Barbosa Ribeiro de Oliveira", "20230012814"],
            ["João Vitor Sampaio Costa", "20230089776"],
            ["Nathan Meira Nóbrega", "20240008904"],
        ],
    )
    doc.add_paragraph()

    # =============== O QUE FOI IMPLEMENTADO ===============
    add_heading(doc, "O que foi implementado", 1)

    # 1. A linguagem Cmd
    add_heading(doc, "1. A linguagem Cmd: comandos e Turing-completude", 2)
    add_body(
        doc,
        "Os compiladores das atividades anteriores traduziam apenas "
        "expressões — sem capacidade de tomar decisões ou repetir "
        "código, o que os deixava aquém de uma linguagem de "
        "programação de verdade. Cmd estende EV (Atividade 08) com "
        "três comandos — condicional (`if`/`else`), repetição "
        "(`while`) e atribuição — e três operadores de comparação "
        "(`<`, `>`, `==`), tornando-a a primeira linguagem "
        "Turing-completa deste conjunto de compiladores.",
    )
    add_body(
        doc,
        "Um programa Cmd é `<decl>* '{' <cmd>* 'return' <exp> ';' '}'`: "
        "zero ou mais declarações (iguais às de EV), seguidas de um "
        "corpo entre chaves com zero ou mais comandos e uma expressão "
        "final obrigatória. `return` não é um comando — só pode "
        "aparecer uma vez, ao final do corpo do programa, nunca dentro "
        "de um `if` ou `while`. A linguagem não tem tipo booleano "
        "separado: `0` é falso, qualquer outro valor é verdadeiro — a "
        "mesma convenção usada nos flags do processador que a geração "
        "de código explora diretamente.",
    )

    # 2. Alteracoes lexicas
    add_heading(doc, "2. Alterações léxicas (lexer.py)", 2)
    add_body(
        doc,
        "Tokens novos: `CHAVE_ESQ`/`CHAVE_DIR` (`{`/`}`), `MENOR`/"
        "`MAIOR`/`IGUAL_IGUAL` (`<`/`>`/`==`), e as palavras-chave "
        "`if`/`else`/`while`/`return`.",
    )
    add_body(doc, "Dois detalhes exigidos pela seção 3 do enunciado:")
    add_bullet(
        doc,
        "`==` vs `=` — ao encontrar `=`, o lexer olha o caractere "
        "seguinte (sem avançar o cursor até decidir); se for outro "
        "`=`, o token é `IGUAL_IGUAL` (consumindo os dois caracteres); "
        "senão, é `IGUAL` (consumindo só um). Implementado em "
        "`_ler_igual()`.",
    )
    add_bullet(
        doc,
        "Palavras-chave — reconhecidas com a mesma regra léxica de "
        "identificador (`letra letra_digito*`) em "
        "`_ler_identificador_ou_palavra_chave()`; só depois de montar "
        "o lexema completo é que ele é comparado contra um dicionário "
        "fixo (`PALAVRAS_CHAVE`) para decidir se vira um token de "
        "palavra-chave ou um `IDENT` comum — exatamente a técnica "
        "sugerida no enunciado. Isso garante que `ifx` continue sendo "
        "um identificador válido, não confundido com a palavra-chave "
        "`if`.",
    )

    # 3. Alteracoes na sintaxe
    add_heading(doc, "3. Alterações na sintaxe (parser.py)", 2)
    add_body(
        doc,
        "A análise do programa começa de forma similar a EV: zero ou "
        "mais declarações (enquanto o próximo token for `IDENT`), "
        "terminadas por `{`. Dentro do bloco, `_analisa_cmd()` decide "
        "por peek de 1 token entre três alternativas — `if`/`while`/"
        "identificador — exatamente como a seção 4 do enunciado "
        "descreve. A lista de comandos de qualquer bloco (programa, "
        "ambos os braços do `if`, corpo do `while`) termina quando o "
        "próximo token não é mais um desses três indicadores.",
    )
    add_body(
        doc,
        "A análise de expressões ganha um nível de precedência novo, "
        "`exp` (comparação), que fica abaixo de `exp_a` (aditivo) na "
        "gramática — os operadores de comparação têm a precedência "
        "mais baixa da linguagem. `_analisa_exp()` segue exatamente o "
        "mesmo formato de laço já usado em `_analisa_exp_a()` e "
        "`_analisa_exp_m()` desde a Atividade 07, só trocando o "
        "conjunto de operadores e a função de nível inferior chamada.",
    )

    # 4. AST
    add_heading(doc, "4. Árvore de sintaxe abstrata (ast_cmd.py)", 2)
    add_body(
        doc,
        "Reaproveita `Exp`/`Const`/`Var`/`OpBin` das atividades "
        "anteriores. `Op` ganha três valores para os operadores de "
        "comparação (`MENOR`, `MAIOR`, `IGUAL`) — `OpBin.avaliar()` "
        "foi estendido com os três casos, devolvendo `1`/`0`.",
    )
    add_body(
        doc,
        "Novos nós de comando, todos `Cmd` (nova classe-base, análoga "
        "a `Exp` mas para comandos):",
    )
    add_bullet(
        doc,
        "`Atrib(nome, exp, posicao)` — igual em espírito ao "
        "`Var.posicao` da Atividade 08: guarda a posição para "
        "mensagens de erro semântico sem afetar a igualdade "
        "estrutural (`field(compare=False)`).",
    )
    add_bullet(
        doc,
        "`If(cond, cmds_then, cmds_else)` — o braço `else` é "
        "obrigatório na gramática, mas `cmds_else` pode ser uma tupla "
        "vazia.",
    )
    add_bullet(doc, "`While(cond, cmds)`.")
    add_bullet(
        doc,
        "`Programa(declaracoes, comandos, exp_final)` — ganhou o "
        "campo `comandos`.",
    )
    add_body(
        doc,
        "`Programa.avaliar()` continua sendo o interpretador de "
        "referência. Para executar os comandos, uma função livre "
        "`_executar(comandos, env)` percorre a sequência mutando o "
        "ambiente: `Atrib` atualiza `env` diretamente; `If` decide "
        "qual braço executar chamando `_executar` recursivamente; "
        "`While` repete a chamada recursiva enquanto a condição for "
        "verdadeira. Essa função serve como o “modelo de execução” "
        "contra o qual validamos a geração de código.",
    )

    # 5. Semantica
    add_heading(doc, "5. Análise semântica (semantica.py)", 2)
    add_body(
        doc,
        "Estende `verifica_programa()` da Atividade 08 com "
        "`_verifica_cmd()`, que percorre recursivamente `If` (condição "
        "+ os dois blocos) e `While` (condição + o bloco). A "
        "verificação nova, exigida pela seção 5 do enunciado, é a de "
        "`Atrib`: um comando de atribuição tem dois componentes a "
        "checar — a expressão do lado direito (não pode referenciar "
        "variável não declarada) e a variável do lado esquerdo "
        "(também precisa já estar declarada). Verificamos primeiro o "
        "lado direito (ordem natural de execução: calcula-se o valor "
        "antes de armazenar), depois o lado esquerdo. Como a "
        "atribuição não cria variáveis novas, nada é inserido na "
        "tabela de símbolos ao processar um `Atrib` — só `Decl` "
        "insere símbolos, exatamente como o enunciado especifica.",
    )

    # 6. Codegen
    add_heading(doc, "6. Geração de código (codegen.py)", 2)
    add_heading(doc, "Comparações (seção 6.1 do enunciado)", 2)
    add_body(
        doc,
        "Reaproveitamos o esquema de pilha já usado para `OpBin` "
        "aritmético (`dir` primeiro, `push`, `esq`, `pop %rbx` → "
        "`%rax = esq`, `%rbx = dir`). Esse esquema já deixa os "
        "operandos exatamente na ordem do exemplo do enunciado para "
        "`A == B` (`%rax = A`, `%rbx = B`), então a comparação foi "
        "encaixada sem nenhuma adaptação:",
    )
    add_code_block(
        doc,
        "xor %rcx, %rcx\n"
        "cmp %rbx, %rax\n"
        "set<cc> %cl\n"
        "mov %rcx, %rax",
    )
    add_body(
        doc,
        "onde `set<cc>` é `setz` (`==`), `setl` (`<`) ou `setg` "
        "(`>`). `%rcx` é usado como temporário porque `SETcc` só "
        "aceita um operando de 8 bits — não dá para usar `%rax` "
        "diretamente.",
    )
    add_heading(doc, "if/else e while (seção 6.2 do enunciado)", 2)
    add_body(
        doc,
        "`GeradorDeCodigo` mantém um contador interno "
        "(`_contador_rotulos`) incrementado a cada `if` ou `while` "
        "gerado, produzindo rótulos únicos (`Lfalso0`/`Lfim0` para o "
        "primeiro `if`, `Linicio1`/`Lfim1` para o próximo `while`, "
        "etc.). Os modelos de tradução seguem exatamente o enunciado:",
    )
    add_code_block(
        doc,
        "<codigo_E>\n"
        "cmp $0, %rax\n"
        "jz LfalsoN\n"
        "<codigo_C1>\n"
        "jmp LfimN\n"
        "LfalsoN:\n"
        "<codigo_C2>\n"
        "LfimN:",
    )
    add_code_block(
        doc,
        "LinicioN:\n"
        "<codigo_E>\n"
        "cmp $0, %rax\n"
        "jz LfimN\n"
        "<codigo_C>\n"
        "jmp LinicioN\n"
        "LfimN:",
    )
    add_body(
        doc,
        "Também adicionamos comentários (`# if cond {`, `# } else {`, "
        "`# }`) delimitando os blocos no `.s` gerado, na mesma linha "
        "da prática já adotada na Atividade 08 para "
        "declarações/atribuições — útil para conferir visualmente que "
        "a estrutura de rótulos e desvios está correta.",
    )
    add_body(
        doc,
        "Modelo de saída: mesmo modelo com seções `.bss`/`.text` da "
        "Atividade 08. As variáveis continuam vindo só das "
        "declarações (`_coleta_variaveis`, deduplicado); o corpo agora "
        "inclui, em ordem: código das declarações, código dos "
        "comandos (que pode conter `if`/`while` com seus próprios "
        "rótulos), e por fim o código da expressão final.",
    )

    # 7. CLI
    add_heading(doc, "7. Ponto de entrada (compcmd.py)", 2)
    add_body(
        doc,
        "Mesma estrutura das atividades anteriores: lê o arquivo "
        "`.cmd` → `analisar()` (léxico + sintático) → "
        "`verifica_programa()` (semântico) → `gerar_programa()` "
        "(codegen) → grava `.s`. Qualquer uma das três exceções é "
        "capturada, imprime a mensagem em `stderr` e encerra com exit "
        "code 1, sem gravar `.s`.",
    )

    # 8. Variacao sintatica
    add_heading(doc, "8. Variação sintática", 2)
    add_body(
        doc,
        "Seguimos exatamente a sintaxe do enunciado, sem adotar "
        "nenhuma das extensões da seção 7: operadores `<=`/`>=`, "
        "operadores booleanos (`AND`/`OR`/`NOT`), comando de entrada, "
        "`if` sem `else`, ou atribuição que cria variáveis novas — "
        "todas apresentadas explicitamente como possibilidades "
        "opcionais, não como parte do artefato mínimo pedido na "
        "seção 8.",
    )

    # 9. Testes
    add_heading(doc, "9. Suíte de testes (tests/test_cmd.py)", 2)
    add_body(doc, "35 testes, 0 falhas, em 7 classes:")
    add_table(
        doc,
        headers=["Classe", "Foco", "Casos"],
        rows=[
            [
                "TestLexico",
                "Tokens novos, == vs =, palavra-chave vs. identificador",
                "6",
            ],
            [
                "TestParser",
                "Estrutura da AST (If/While/Atrib, precedência da comparação)",
                "6",
            ],
            [
                "TestErrosSintaticos",
                "Entradas mal formadas (if sem else, bloco sem chaves, etc.)",
                "5",
            ],
            [
                "TestSemantica",
                "Atribuição a variável não declarada (ambos os lados, inclusive aninhada)",
                "5",
            ],
            [
                "TestInterpretacao",
                "Programa.avaliar() para os 4 exemplos do enunciado",
                "5",
            ],
            [
                "TestEquivalenciaSemantica",
                "Código gerado bate com o interpretador (10 programas, incluindo laços)",
                "1",
            ],
            [
                "TestCodegen",
                "Rótulos únicos por if/while, modelo completo",
                "3",
            ],
            [
                "TestCLI",
                "Comportamento de compcmd.py como subprocesso",
                "4",
            ],
            ["Total", "", "35"],
        ],
    )
    doc.add_paragraph()
    add_body(doc, "Destaques:")
    add_bullet(
        doc,
        "`test_discriminante_do_enunciado`, `test_soma_do_enunciado`, "
        "`test_resto_da_divisao` e `test_mdc` reproduzem os quatro "
        "exemplos do enunciado (seções 2 e 9) e confirmam os valores "
        "citados: `8`, `45`, `2` e `6`, respectivamente.",
    )
    add_bullet(
        doc,
        "`TestEquivalenciaSemantica` é o teste mais importante desta "
        "atividade: o simulador de máquina de pilha das Atividades "
        "06-08 foi estendido com um program counter e um mapa de "
        "rótulos, capaz de executar `jmp`/`jz` de verdade — inclusive "
        "laços (`while`), com um limite de passos (`MAX_PASSOS`) para "
        "detectar loop infinito por bug em vez de travar o teste. Isso "
        "prova que o código gerado para `if`/`while`/comparações é "
        "semanticamente equivalente ao interpretador, incluindo os "
        "quatro programas com `while` (soma, resto, mdc — que tem "
        "laços aninhados).",
    )
    add_bullet(
        doc,
        "`test_rotulos_unicos_para_ifs_diferentes` confirma que dois "
        "`if` no mesmo programa geram pares de rótulos distintos "
        "(`Lfalso0`/`Lfim0` e `Lfalso1`/`Lfim1`), evitando colisão de "
        "símbolos no assembly.",
    )

    # 10. Exemplos
    add_heading(doc, "10. Arquivos de exemplo", 2)
    add_body(
        doc,
        "Quatro programas válidos — os dois exemplos da seção 2 do "
        "enunciado (discriminante e soma) e os dois exemplos "
        "adicionais da seção 9 (resto da divisão e MDC, ambos usando "
        "`while` e o truque de `m + 1 > n` para simular “maior ou "
        "igual”, já que Cmd não tem esse operador) — e dois "
        "inválidos: atribuição a uma variável nunca declarada, e "
        "atribuição cuja expressão do lado direito referencia uma "
        "variável não declarada dentro de um `if`.",
    )

    # =============== EXEMPLO DE SAIDA ===============
    add_heading(doc, "Exemplo de saída — discriminante (trecho)", 1)
    add_code_block(
        doc,
        "# a = 1;\n"
        "mov $1, %rax\n"
        "mov %rax, a\n"
        "...\n"
        "# delta = ((b * b) - ((4 * a) * c));\n"
        "...\n"
        "# if (delta < 0) {\n"
        "mov $0, %rax\n"
        "push %rax\n"
        "mov delta, %rax\n"
        "pop %rbx\n"
        "xor %rcx, %rcx\n"
        "cmp %rbx, %rax\n"
        "setl %cl\n"
        "mov %rcx, %rax\n"
        "cmp $0, %rax\n"
        "jz Lfalso0\n"
        "# delta = (0 - delta);\n"
        "...\n"
        "jmp Lfim0\n"
        "Lfalso0:\n"
        "# } else {\n"
        "# delta = delta;\n"
        "...\n"
        "Lfim0:\n"
        "# }\n"
        "# return delta;\n"
        "mov delta, %rax\n"
        "call imprime_num\n"
        "call sair",
    )
    add_body(
        doc,
        "Executado, imprime `8` (delta calculado é `-8`, e o `if` "
        "inverte o sinal porque `delta < 0`).",
    )

    # =============== ESTRUTURA DE ARQUIVOS ===============
    add_heading(doc, "Estrutura de arquivos entregue", 1)
    add_code_block(
        doc,
        "compilador-cmd/\n"
        "├── lexer.py\n"
        "├── ast_cmd.py\n"
        "├── parser.py\n"
        "├── semantica.py\n"
        "├── codegen.py\n"
        "├── compcmd.py\n"
        "├── runtime.s\n"
        "├── exemplos/\n"
        "│   ├── valido1.cmd\n"
        "│   ├── valido2.cmd\n"
        "│   ├── valido3.cmd\n"
        "│   ├── valido4.cmd\n"
        "│   ├── invalido_atrib_var_nao_declarada.cmd\n"
        "│   └── invalido_atrib_direita_nao_declarada.cmd\n"
        "├── tests/\n"
        "│   └── test_cmd.py\n"
        "├── README.md\n"
        "├── PLANO.md\n"
        "└── RELATORIO.md",
    )

    # =============== DECISOES DE PROJETO ===============
    add_heading(doc, "Decisões de projeto", 1)

    add_heading(
        doc,
        "Por que os operandos da comparação já saem na ordem certa "
        "do esquema de pilha existente?",
        2,
    )
    add_body(
        doc,
        "Porque o esquema de pilha das atividades anteriores (`dir` "
        "primeiro, `push`, `esq`, `pop %rbx`) sempre deixa `%rax = "
        "esq` e `%rbx = dir` depois do `pop`. O exemplo do enunciado "
        "para `A == B` assume exatamente essa disposição (`%rax = A`, "
        "`%rbx = B`) antes do `cmp`. Não precisamos adaptar nada — a "
        "mesma técnica de tradução usada para +/-/*// desde a "
        "Atividade 06 já resolve comparações de graça.",
    )

    add_heading(
        doc,
        "Por que usar um contador de instância (_contador_rotulos) em "
        "vez de nomes de rótulo determinados pela posição do if/while "
        "na árvore?",
        2,
    )
    add_body(
        doc,
        "É a técnica sugerida literalmente pelo enunciado (“manter um "
        "contador de quantos rótulos foram gerados até agora”) e é a "
        "mais simples que garante unicidade sem exigir nenhuma "
        "análise adicional da árvore (como profundidade de "
        "aninhamento ou um caminho único por nó). Cada if/while, não "
        "importa o quão aninhado, recebe um número novo.",
    )

    add_heading(
        doc,
        "Por que o simulador de teste precisa de um program counter "
        "explícito, e não apenas percorrer as linhas em ordem?",
        2,
    )
    add_body(
        doc,
        "Porque agora o código gerado tem desvios de verdade "
        "(`jmp`/`jz`) que podem andar para trás no texto (voltar para "
        "`Linicio0:` em um `while`), o que o simulador linear das "
        "atividades anteriores (que só executava de cima para baixo) "
        "não suportava. Construímos um mapa `rotulo → índice` e um "
        "laço com `pc` que só avança sequencialmente quando a "
        "instrução não é um desvio — exatamente como um processador "
        "de verdade interpretaria o mesmo código.",
    )

    add_heading(
        doc,
        "Por que checar o lado direito da atribuição antes do lado "
        "esquerdo na análise semântica?",
        2,
    )
    add_body(
        doc,
        "Reflete a ordem natural de execução do comando: primeiro se "
        "calcula o valor da expressão, só depois ele é armazenado na "
        "variável. O enunciado não impõe uma ordem específica entre "
        "as duas checagens (ambas resultam em erro, de qualquer "
        "forma), mas essa ordem é a mais intuitiva e é a que usamos "
        "consistentemente.",
    )

    add_heading(
        doc,
        "Por que os comentários de bloco (# if cond {, # } else {, "
        "# }) no código gerado?",
        2,
    )
    add_body(
        doc,
        "Mesma justificativa da Atividade 08: custo desprezível, e "
        "tornam o `.s` muito mais fácil de inspecionar visualmente — "
        "foi assim que confirmamos rapidamente, durante o "
        "desenvolvimento, que a estrutura de rótulos de um `if` dentro "
        "de outro `if` estava correta antes mesmo de rodar os testes "
        "automatizados.",
    )

    # =============== DIFICULDADES ===============
    add_heading(doc, "Dificuldades", 1)
    add_body(
        doc,
        "Nenhuma dificuldade significativa. O enunciado é bastante "
        "explícito sobre os modelos de tradução de comparações, `if` "
        "e `while` (inclusive com o assembly completo de exemplo), o "
        "que tornou a implementação do codegen bem direta. Os pontos "
        "que exigiram mais atenção:",
    )
    add_bullet(
        doc,
        "Escrever o simulador de teste com desvios reais. Diferente "
        "das atividades anteriores, onde o “simulador” era apenas um "
        "interpretador linear de cima para baixo, aqui foi necessário "
        "simular um program counter de verdade, com salto para trás "
        "(loop) e um limite de passos de segurança para não travar a "
        "suíte de testes em caso de bug que gerasse um loop infinito "
        "de verdade.",
    )
    add_bullet(
        doc,
        "Diferenciar `==` de `=` no lexer. Resolvido com um lookahead "
        "explícito de 1 caractere antes de decidir o tipo do token, "
        "sem consumir o segundo `=` a menos que ele realmente esteja "
        "lá.",
    )
    add_bullet(
        doc,
        "Garantir que `ifx` não vire a palavra-chave `if`. A regra de "
        "reconhecer o identificador inteiro primeiro e só depois "
        "comparar contra a tabela de palavras-chave resolve isso "
        "automaticamente, mas foi importante escrever um teste "
        "específico para confirmar.",
    )

    doc.save(path)


if __name__ == "__main__":
    import os
    out = os.path.join(os.path.dirname(__file__), "RELATORIO.docx")
    build_document(out)
    print(f"gerado: {out}")
