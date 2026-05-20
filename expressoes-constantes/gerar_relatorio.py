"""Gera o arquivo RELATORIO.docx a partir do conteudo do RELATORIO.md.

Uso:
    python gerar_relatorio.py

Requer: python-docx (pip install python-docx).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT = "Calibri"
BODY_SIZE = Pt(11)
CODE_FONT = "Consolas"
HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x5F)
TABLE_HEADER_BG = "1F3A5F"
CODE_BG = "F4F4F4"


# -------- helpers --------

def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_paragraph_shading(paragraph, hex_color):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.bold = True
    if level == 0:
        run.font.size = Pt(20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = HEADING_COLOR
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = HEADING_COLOR
    else:
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(14 if level > 0 else 0)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_runs_with_inline_code(p, text):
    """Quebra o texto por trechos entre crases e adiciona como runs separados."""
    parts = text.split("`")
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.font.name = FONT
        run.font.size = BODY_SIZE
        if i % 2 == 1:
            run.font.name = CODE_FONT
            run.font.size = Pt(10)


def add_body(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(6)
    add_runs_with_inline_code(p, text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_runs_with_inline_code(p, text)
    return p


def add_code_block(doc, code):
    """Bloco de codigo pre-formatado: monoespacada, fundo cinza, sem quebra de linha forcada."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    set_paragraph_shading(p, CODE_BG)
    linhas = code.rstrip("\n").split("\n")
    for i, linha in enumerate(linhas):
        if i > 0:
            run = p.add_run()
            run.add_break()
        run = p.add_run(linha)
        run.font.name = CODE_FONT
        run.font.size = Pt(9.5)
    return p


def add_centered_info_line(doc, label, value, bold_label=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    if label:
        r = p.add_run(label + " ")
        r.bold = bold_label
        r.font.name = FONT
        r.font.size = BODY_SIZE
    r2 = p.add_run(value)
    r2.font.name = FONT
    r2.font.size = BODY_SIZE


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for idx, h in enumerate(headers):
        hdr_cells[idx].text = ""
        r = hdr_cells[idx].paragraphs[0].add_run(h)
        r.bold = True
        r.font.name = FONT
        r.font.size = BODY_SIZE
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[idx], TABLE_HEADER_BG)
        hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if col_widths:
            hdr_cells[idx].width = col_widths[idx]
    for row in rows:
        cells = table.add_row().cells
        for idx, valor in enumerate(row):
            cells[idx].text = ""
            # suporta codigo inline em celulas
            add_runs_with_inline_code(cells[idx].paragraphs[0], str(valor))
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if col_widths:
                cells[idx].width = col_widths[idx]
    return table


# -------- documento --------

def build_document(path):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = BODY_SIZE

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # Cabecalho institucional
    add_centered_info_line(doc, "", "Universidade Federal da Paraíba (UFPB)")
    add_centered_info_line(doc, "", "Centro de Informática – Curso de Ciência da Computação")
    add_centered_info_line(doc, "Disciplina:", "Construção de Compiladores 1")
    add_centered_info_line(doc, "Professor:", "Andrei de Araújo Formiga")
    doc.add_paragraph()

    # Titulo
    add_heading(doc, "Relatório – Atividade 04", 0)
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Análise Léxica EC1 (Expressões Constantes 1)")
    r_sub.italic = True
    r_sub.font.size = Pt(13)
    r_sub.font.name = FONT
    p_sub.paragraph_format.space_after = Pt(18)

    # Integrantes
    add_heading(doc, "Integrantes do grupo", 1)
    add_table(
        doc,
        headers=["Nome", "Matrícula"],
        rows=[
            ["Davi Alves Rodrigues", "20230102377"],
            ["Gabriel Barbosa Ribeiro de Oliveira", "20230012814"],
            ["João Vitor Sampaio Costa", "20230089776"],
            ["Nathan Meira Nóbrega", "20240008904"],
        ],
    )
    doc.add_paragraph()

    # =============== O QUE FOI IMPLEMENTADO ===============
    add_heading(doc, "O que foi implementado", 1)

    # 1. Estrutura do token
    add_heading(doc, "1. Estrutura do token (Token, TipoToken)", 2)
    add_body(doc, "Foi definida a enumeração `TipoToken` com oito valores:")
    add_table(
        doc,
        headers=["Enum", "Descrição"],
        rows=[
            ["NUMERO", "Sequência de dígitos"],
            ["PAREN_ESQ", "("],
            ["PAREN_DIR", ")"],
            ["SOMA", "+"],
            ["SUB", "-"],
            ["MULT", "*"],
            ["DIV", "/"],
            ["EOF", "Fim da entrada"],
        ],
    )
    doc.add_paragraph()
    add_body(
        doc,
        "A classe `Token` é um `dataclass` Python com três campos: `tipo`, `lexema` e "
        "`posicao`. O método `__str__` produz exatamente o formato exigido pelo enunciado:",
    )
    add_code_block(doc, '<ParenEsq, "(", 0>\n<Numero, "33", 1>')
    add_body(
        doc,
        "A saída para o exemplo do enunciado `(33 + (912 * 11))` bate caractere a "
        "caractere com o esperado, incluindo as posições:",
    )
    add_code_block(
        doc,
        '<ParenEsq, "(", 0>    <Numero, "33", 1>    <Soma, "+", 4>\n'
        '<ParenEsq, "(", 6>    <Numero, "912", 7>   <Mult, "*", 11>\n'
        '<Numero, "11", 13>    <ParenDir, ")", 15>  <ParenDir, ")", 16>',
    )

    # 2. Tratamento de erros lexicos
    add_heading(doc, "2. Tratamento de erros léxicos (ErroLexico)", 2)
    add_body(
        doc,
        "A exceção `ErroLexico` é levantada assim que um caractere inválido é "
        "encontrado. Ela armazena a posição do erro e o caractere responsável, "
        "produzindo uma mensagem como:",
    )
    add_code_block(doc, "Erro léxico na posição 4: caractere inesperado 'x' (ASCII 120)")
    add_body(
        doc,
        "O processo termina com código de saída `1`, separando a mensagem de erro "
        "para `stderr` e a sequência de tokens para `stdout`, conforme boa prática "
        "de CLIs.",
    )

    # 3. Analisador lexico
    add_heading(doc, "3. Analisador léxico (AnalisadorLexico)", 2)
    add_body(
        doc,
        "A classe `AnalisadorLexico` mantém internamente um ponteiro de posição "
        "(`_pos`) sobre a string de entrada. Duas interfaces públicas foram "
        "implementadas, conforme o enunciado descreve as duas opções válidas:",
    )
    add_bullet(
        doc,
        "`proximo_token()` — retorna um token por vez; retorna `EOF` ao esgotar a "
        "entrada.",
    )
    add_bullet(doc, "`tokenizar()` — retorna todos os tokens em lista (sem EOF).")
    add_body(
        doc,
        "Descarte de espaços em branco: os quatro tipos citados no enunciado são "
        "ignorados — espaço (`0x20`), tabulação (`0x09`), nova linha (`0x0A`) e "
        "retorno do carro (`0x0D`).",
    )
    add_body(
        doc,
        "Agrupamento de números: ao encontrar um dígito, o lexer avança enquanto "
        "houver dígitos consecutivos, reunindo todos em um único token `NUMERO` "
        "com o lexema completo.",
    )
    add_body(
        doc,
        "Caracteres simples: parênteses e operadores são reconhecidos com lookup "
        "direto em dicionário, resultando em O(1) por token.",
    )

    # 4. Extensao opcional
    add_heading(doc, "4. Extensão opcional: comentários de linha (#)", 2)
    add_body(
        doc,
        "O enunciado menciona explicitamente que os grupos podem adicionar suporte "
        "a comentários. Foi implementado o comentário de linha iniciado por `#`: "
        "tudo da `#` até o fim da linha (inclusive a ausência de `\\n` no final do "
        "arquivo) é descartado silenciosamente.",
    )
    add_body(doc, "Exemplo de uso:")
    add_code_block(doc, "# calcula (6 * 7)\n(6 * 7)   # resultado: 42")

    # 5. Ponto de entrada
    add_heading(doc, "5. Ponto de entrada (main)", 2)
    add_body(
        doc,
        "O script aceita um argumento de linha de comando com o caminho do arquivo "
        "de entrada:",
    )
    add_code_block(doc, "python lexer.py exemplos/valido1.ec1")
    add_body(
        doc,
        "Erros de arquivo não encontrado e erros léxicos são reportados no `stderr` "
        "com código de saída `1`. A sequência de tokens vai para o `stdout`, "
        "facilitando redirecionamento e uso em pipelines de testes.",
    )

    # 6. Suite de testes
    add_heading(doc, "6. Suite de testes (43 casos)", 2)
    add_body(
        doc,
        "Todos os testes foram escritos com o módulo `unittest` da biblioteca padrão "
        "do Python, sem dependências externas. Execução:",
    )
    add_code_block(doc, "python tests/test_lexer.py")
    add_body(doc, "Resultado: 43 testes, 0 falhas.")
    add_body(doc, "Distribuição por categoria:")
    add_table(
        doc,
        headers=["Categoria", "Casos"],
        rows=[
            ["Números (tipo, lexema, posição)", "5"],
            ["Operadores (+, -, *, /)", "4"],
            ["Parênteses", "3"],
            ["Expressões completas", "6"],
            ["Espaços em branco", "8"],
            ["Comentários (extensão #)", "4"],
            ["Erros léxicos", "7"],
            ["Token EOF", "2"],
            ["Representação str", "3"],
            ["Total", "43"],
        ],
    )
    doc.add_paragraph()
    add_body(doc, "Destaques:")
    add_bullet(
        doc,
        "`test_exemplo_do_enunciado` verifica token por token (tipo, lexema e "
        "posição) contra o exemplo literal do enunciado, garantindo conformidade "
        "total.",
    )
    add_bullet(
        doc,
        "`test_posicao_com_espacos` verifica que as posições são do caractere de "
        "início do token na string original, não do token na sequência de tokens.",
    )
    add_bullet(
        doc,
        "Todos os testes de erro verificam não apenas que o `ErroLexico` é "
        "levantado, mas também que a posição reportada é a correta.",
    )

    # 7. Arquivos de exemplo
    add_heading(doc, "7. Arquivos de exemplo", 2)
    add_body(doc, "Quatro arquivos `.ec1` foram criados em `exemplos/`:")
    add_table(
        doc,
        headers=["Arquivo", "Conteúdo", "Resultado esperado"],
        rows=[
            ["valido1.ec1", "(33 + (912 * 11))", "9 tokens válidos"],
            ["valido2.ec1", "((427 / 7) + (11 * (231 + 5)))", "17 tokens válidos"],
            [
                "valido3.ec1",
                "Expressão aninhada com comentário e indentação",
                "13 tokens válidos",
            ],
            ["invalido.ec1", "(12 x 5)", "Erro léxico na posição 4"],
        ],
    )

    # =============== ESTRUTURA DE ARQUIVOS ===============
    add_heading(doc, "Estrutura de arquivos entregue", 1)
    add_code_block(
        doc,
        "ec1_lexer/\n"
        "├── lexer.py              ← implementação principal\n"
        "├── tests/\n"
        "│   └── test_lexer.py     ← 43 testes automatizados\n"
        "├── exemplos/\n"
        "│   ├── valido1.ec1\n"
        "│   ├── valido2.ec1\n"
        "│   ├── valido3.ec1\n"
        "│   └── invalido.ec1\n"
        "├── README.md             ← instruções de uso e execução\n"
        "├── PLANO.md              ← plano das próximas etapas\n"
        "└── RELATORIO.md          ← este arquivo",
    )

    # =============== DECISOES DE PROJETO ===============
    add_heading(doc, "Decisões de projeto", 1)

    add_heading(doc, "Por que dataclass para Token?", 2)
    add_body(
        doc,
        "Fornece `__eq__` automático por valor, o que permite comparações diretas "
        "nos testes (`self.assertEqual(toks, esperado)`) sem código extra.",
    )

    add_heading(
        doc,
        "Por que separar cada operador em tipo próprio em vez de um tipo genérico OPERADOR?",
        2,
    )
    add_body(
        doc,
        "O enunciado recomenda explicitamente esta abordagem para facilitar as "
        "etapas seguintes do compilador (análise sintática e geração de código), "
        "que precisarão distinguir os operadores.",
    )

    add_heading(doc, "Por que frozenset para ESPACOS e dict para CHAR_SIMPLES?", 2)
    add_body(
        doc,
        "Ambos oferecem lookup O(1). O `frozenset` deixa claro que o conjunto é "
        "imutável; o `dict` mapeia diretamente do caractere para o tipo, eliminando "
        "um bloco de `if/elif`.",
    )

    add_heading(doc, "Tratamento de erro: parar no primeiro erro vs. continuar", 2)
    add_body(
        doc,
        "O enunciado apresenta as duas opções. Foi escolhida a abordagem de parar "
        "no primeiro erro, que é mais simples de implementar e suficiente para "
        "esta etapa. A interface `proximo_token()` levanta `ErroLexico` "
        "imediatamente, o que permite que um eventual analisador sintático futuro "
        "decida como tratar a situação.",
    )

    doc.save(path)


if __name__ == "__main__":
    import os
    out = os.path.join(os.path.dirname(__file__), "RELATORIO.docx")
    build_document(out)
    print(f"gerado: {out}")
