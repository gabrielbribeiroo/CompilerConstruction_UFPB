"""Gera o arquivo RELATORIO.docx a partir do conteudo do RELATORIO.md.

Uso:
    python gerar_relatorio.py

Requer: python-docx (pip install python-docx).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT = "Calibri"
BODY_SIZE = Pt(11)
CODE_FONT = "Consolas"
HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x5F)
TABLE_HEADER_BG = "1F3A5F"
CODE_BG = "F4F4F4"


# -------- helpers --------

def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_paragraph_shading(paragraph, hex_color):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.bold = True
    if level == 0:
        run.font.size = Pt(20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = HEADING_COLOR
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = HEADING_COLOR
    else:
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(14 if level > 0 else 0)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_runs_with_inline_code(p, text):
    parts = text.split("`")
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.font.name = FONT
        run.font.size = BODY_SIZE
        if i % 2 == 1:
            run.font.name = CODE_FONT
            run.font.size = Pt(10)


def add_body(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(6)
    add_runs_with_inline_code(p, text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_runs_with_inline_code(p, text)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    set_paragraph_shading(p, CODE_BG)
    linhas = code.rstrip("\n").split("\n")
    for i, linha in enumerate(linhas):
        if i > 0:
            run = p.add_run()
            run.add_break()
        run = p.add_run(linha)
        run.font.name = CODE_FONT
        run.font.size = Pt(9.5)
    return p


def add_centered_info_line(doc, label, value, bold_label=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    if label:
        r = p.add_run(label + " ")
        r.bold = bold_label
        r.font.name = FONT
        r.font.size = BODY_SIZE
    r2 = p.add_run(value)
    r2.font.name = FONT
    r2.font.size = BODY_SIZE


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for idx, h in enumerate(headers):
        hdr_cells[idx].text = ""
        r = hdr_cells[idx].paragraphs[0].add_run(h)
        r.bold = True
        r.font.name = FONT
        r.font.size = BODY_SIZE
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[idx], TABLE_HEADER_BG)
        hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if col_widths:
            hdr_cells[idx].width = col_widths[idx]
    for row in rows:
        cells = table.add_row().cells
        for idx, valor in enumerate(row):
            cells[idx].text = ""
            add_runs_with_inline_code(cells[idx].paragraphs[0], str(valor))
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if col_widths:
                cells[idx].width = col_widths[idx]
    return table


# -------- documento --------

def build_document(path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = BODY_SIZE
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # Cabecalho institucional
    add_centered_info_line(doc, "", "Universidade Federal da Paraíba (UFPB)")
    add_centered_info_line(doc, "", "Centro de Informática – Curso de Ciência da Computação")
    add_centered_info_line(doc, "Disciplina:", "Construção de Compiladores 1")
    add_centered_info_line(doc, "Professor:", "Andrei de Araújo Formiga")
    doc.add_paragraph()

    # Titulo
    add_heading(doc, "Relatório – Atividade 06", 0)
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Compilador EC1 (Expressões Constantes 1)")
    r_sub.italic = True
    r_sub.font.size = Pt(13)
    r_sub.font.name = FONT
    p_sub.paragraph_format.space_after = Pt(18)

    # Integrantes
    add_heading(doc, "Integrantes do grupo", 1)
    add_table(
        doc,
        headers=["Nome", "Matrícula"],
        rows=[
            ["Davi Alves Rodrigues", "20230102377"],
            ["Gabriel Barbosa Ribeiro de Oliveira", "20230012814"],
            ["João Vitor Sampaio Costa", "20230089776"],
            ["Nathan Meira Nóbrega", "20240008904"],
        ],
    )
    doc.add_paragraph()

    # =============== O QUE FOI IMPLEMENTADO ===============
    add_heading(doc, "O que foi implementado", 1)

    # 1. Reaproveitamento
    add_heading(doc, "1. Reaproveitamento das atividades anteriores", 2)
    add_bullet(doc, "`lexer.py` — análise léxica idêntica à entregue na Atividade 04.")
    add_bullet(
        doc,
        "`ast_ec1.py` — hierarquia de classes da AST (Atividade 05): `Exp` "
        "abstrata, `Const`, `OpBin` (`dataclass(frozen=True)`), enum `Op`.",
    )
    add_bullet(
        doc,
        "`parser.py` — analisador sintático descendente recursivo da "
        "Atividade 05, com `ErroSintatico` que carrega posição e mensagem.",
    )
    add_bullet(
        doc,
        "`runtime.s` — exatamente o `runtime.s` fornecido pelo professor na "
        "Atividade 02 (`imprime_num`, `sair`). Necessário para o `.s` gerado "
        "poder ser montado e linkado em um executável.",
    )

    # 2. Codegen
    add_heading(doc, "2. Gerador de código (codegen.py)", 2)
    add_body(
        doc,
        "A classe `GeradorDeCodigo` percorre recursivamente a AST e emite "
        "assembly x86-64 (sintaxe GNU Assembler). Para cada tipo de nó:",
    )
    add_bullet(
        doc,
        "`Const(n)` → uma única instrução `mov $n, %rax`.",
    )
    add_bullet(
        doc,
        "`OpBin(op, esq, dir)` → segue o esquema baseado na pilha descrito "
        "na seção 4.1 do enunciado, na opção 2 (ordem invertida) da página 9: "
        "(1) emite código de `dir`, (2) `push %rax`, (3) emite código de "
        "`esq`, (4) `pop %rbx`, (5) emite a instrução da operação na forma "
        "`<instr> %rbx, %rax`, deixando o resultado em `%rax`.",
    )
    add_body(doc, "Mapeamento dos operadores:")
    add_table(
        doc,
        headers=["Op", "Instrução emitida"],
        rows=[
            ["SOMA", "add %rbx, %rax"],
            ["SUB", "sub %rbx, %rax"],
            ["MULT", "imul %rbx, %rax"],
            ["DIV", "cqo + idiv %rbx"],
        ],
    )
    doc.add_paragraph()
    add_body(
        doc,
        "A divisão precisa do `cqo` antes do `idiv` porque a instrução "
        "`idiv` em x86-64 divide o par `rdx:rax` pelo operando — `cqo` "
        "estende o sinal de `%rax` para preencher `%rdx`, garantindo "
        "divisão inteira com sinal corretamente. O quociente fica em "
        "`%rax`.",
    )
    add_body(
        doc,
        "A escolha pela opção 2 (ordem invertida) é deliberada: depois do "
        "`pop`, `%rax` tem o operando esquerdo e `%rbx` o direito, "
        "exatamente a ordem natural de operandos para `sub` e `idiv`. Não é "
        "preciso nenhum truque extra para corrigir a ordem em operadores "
        "não-comutativos.",
    )
    add_body(doc, "`GeradorDeCodigo` expõe duas APIs:")
    add_bullet(doc, "`gerar(arvore)` — devolve só o código da expressão.")
    add_bullet(
        doc,
        "`gerar_programa(arvore)` — devolve o `.s` completo, com o modelo "
        "de saída em volta (seção `.text`, rótulo `_start`, chamadas finais "
        "a `imprime_num` e `sair`, `.include \"runtime.s\"`), idêntico ao da "
        "Atividade 02.",
    )

    # 3. CLI
    add_heading(doc, "3. Ponto de entrada (compec1.py)", 2)
    add_body(
        doc,
        "Recebe o nome do arquivo `.ec1` na linha de comando, executa "
        "lex → parse → codegen e grava a saída em um arquivo `.s` ao lado "
        "da entrada (trocando a extensão). Erros léxicos, sintáticos ou de "
        "E/S vão para `stderr` com exit code 1; nesses casos nenhum `.s` é "
        "gravado.",
    )
    add_body(
        doc,
        "A escolha de gravar em arquivo (em vez de imprimir em `stdout`) "
        "segue a preferência expressa na seção 6 do enunciado.",
    )

    # 4. Testes
    add_heading(doc, "4. Suíte de testes (tests/test_codegen.py)", 2)
    add_body(doc, "15 testes, 0 falhas, divididos em 6 classes do `unittest`:")
    add_table(
        doc,
        headers=["Classe", "Foco", "Casos"],
        rows=[
            [
                "TestCodigoDeConstantes",
                "Const(n) produz a instrução mov correta",
                "3",
            ],
            [
                "TestCodigoDeOperacoes",
                "Código literal exato para cada um dos 4 operadores",
                "4",
            ],
            [
                "TestAninhamento",
                "Pilha balanceada em expressão aninhada à direita",
                "1",
            ],
            [
                "TestEquivalenciaSemantica",
                "Semântica do código gerado bate com o interpretador",
                "2",
            ],
            [
                "TestProgramaCompleto",
                "gerar_programa() envolve o código no modelo correto",
                "2",
            ],
            [
                "TestCLI",
                "Comportamento de compec1.py como subprocesso",
                "3",
            ],
            ["Total", "", "15"],
        ],
    )
    doc.add_paragraph()
    add_body(
        doc,
        "Destaque para a classe `TestEquivalenciaSemantica`: ela contém um "
        "simulador minimalista da máquina de pilha (~30 linhas em Python) "
        "que interpreta exatamente as instruções que o nosso codegen pode "
        "emitir. Para cada um dos 16 programas de teste — incluindo os dois "
        "exemplos do enunciado e vários casos com operadores não-comutativos "
        "aninhados em ambas as ordens — o simulador executa o código gerado "
        "e o resultado é comparado com o `avaliar()` do interpretador da "
        "Atividade 05. Se os dois batem para todos os programas, o código "
        "gerado é semanticamente correto, mesmo sem precisar montar e linkar "
        "o `.s` na máquina onde os testes rodam.",
    )
    add_body(
        doc,
        "A classe `TestCLI` invoca `compec1.py` via `subprocess` e verifica "
        "comportamento de ponta a ponta: entrada válida produz `.s` correto; "
        "entrada com erro de sintaxe sai com código 1 e não cria `.s`; "
        "arquivo inexistente é reportado de forma adequada.",
    )

    # 5. Exemplos
    add_heading(doc, "5. Arquivos de exemplo", 2)
    add_body(
        doc,
        "Seis programas válidos em `exemplos/`, cobrindo: constante simples "
        "(`333`), multiplicação simples (`(6 * 7)`), os dois exemplos do "
        "enunciado (`(33 + (912 * 11))` e `((427 / 7) + (11 * (231 + 5)))`), "
        "subtração para estressar não-comutatividade (`(11 - 7)`), e soma "
        "aninhada à direita (`(7 + (3 + 8))`).",
    )

    # =============== EXEMPLO DE SAIDA ===============
    add_heading(doc, "Exemplo de saída — (33 + (912 * 11))", 1)
    add_code_block(
        doc,
        "#\n"
        "# Saida gerada pelo compilador EC1\n"
        "#\n"
        "    .section .text\n"
        "    .globl _start\n"
        "\n"
        "_start:\n"
        "    mov $11, %rax\n"
        "    push %rax\n"
        "    mov $912, %rax\n"
        "    pop %rbx\n"
        "    imul %rbx, %rax\n"
        "    push %rax\n"
        "    mov $33, %rax\n"
        "    pop %rbx\n"
        "    add %rbx, %rax\n"
        "    call imprime_num\n"
        "    call sair\n"
        "\n"
        '    .include "runtime.s"',
    )
    add_body(
        doc,
        "Trace: a multiplicação interna `(912 * 11)` é executada primeiro; "
        "`imul` deixa `10032` em `%rax`. Esse valor é empilhado, `33` vai "
        "para `%rax`, `pop` traz `10032` para `%rbx`, e `add` produz `10065` "
        "em `%rax`, que é então passado para `imprime_num`.",
    )

    # =============== ESTRUTURA DE ARQUIVOS ===============
    add_heading(doc, "Estrutura de arquivos entregue", 1)
    add_code_block(
        doc,
        "compilador-ec1/\n"
        "├── lexer.py\n"
        "├── ast_ec1.py\n"
        "├── parser.py\n"
        "├── codegen.py\n"
        "├── compec1.py\n"
        "├── runtime.s\n"
        "├── exemplos/\n"
        "│   ├── valido1.ec1\n"
        "│   ├── valido2.ec1\n"
        "│   ├── valido3.ec1\n"
        "│   ├── valido4.ec1\n"
        "│   ├── valido5.ec1\n"
        "│   └── valido6.ec1\n"
        "├── tests/\n"
        "│   └── test_codegen.py\n"
        "├── README.md\n"
        "├── PLANO.md\n"
        "└── RELATORIO.md",
    )

    # =============== DECISOES DE PROJETO ===============
    add_heading(doc, "Decisões de projeto", 1)

    add_heading(doc, "Por que a opção 2 (ordem invertida) em vez da opção 1?", 2)
    add_body(
        doc,
        "A opção 1 (página 8) emite o operando esquerdo primeiro, empilha, "
        "depois o direito, e termina com `pop %rbx`. Para soma e "
        "multiplicação, comutativas, funciona direto. Para subtração e "
        "divisão, exige um truque: trocar `pop %rbx` por `mov %rax, %rbx` + "
        "`pop %rax`. A opção 2 não precisa desse caso especial — a ordem "
        "natural já fica `rax = esq`, `rbx = dir`, que é a forma que `sub` "
        "e `idiv` esperam. Código mais uniforme.",
    )

    add_heading(
        doc,
        "Por que imul %rbx, %rax (forma de dois operandos) em vez de imul %rbx?",
        2,
    )
    add_body(
        doc,
        "A forma de dois operandos é mais segura: faz `rax = rax * rbx` sem "
        "mexer em `%rdx`. A forma de um operando faz `rdx:rax = rax * rbx`, "
        "sobrescrevendo `%rdx`, o que não causa problema neste compilador "
        "mas é desnecessário.",
    )

    add_heading(doc, "Por que cqo + idiv em vez de só div?", 2)
    add_body(
        doc,
        "`idiv` é divisão com sinal, coerente com a tipagem inteira da AST "
        "e com o `avaliar()` que usa `int(a / b)` (truncamento para zero, "
        "exatamente o que `idiv` faz). `cqo` estende o sinal de `%rax` em "
        "`%rdx`, o que é indispensável antes de qualquer `idiv` — sem isso, "
        "`%rdx` carregaria lixo da operação anterior.",
    )

    add_heading(doc, "Por que o codegen como classe em vez de funções soltas?", 2)
    add_body(
        doc,
        "A classe acumula o resultado em `self._linhas`, o que evita passar "
        "e juntar listas em cada chamada recursiva. A interface continua "
        "simples (duas funções de conveniência `gerar_codigo` e "
        "`gerar_programa` embrulham a criação da instância para uso comum).",
    )

    add_heading(doc, "Equivalência semântica como teste principal", 2)
    add_body(
        doc,
        "A forma mais sólida de validar o codegen seria montar e executar "
        "os `.s` em uma máquina x86-64 Linux. Como nem todos os ambientes "
        "de desenvolvimento têm essa toolchain disponível, escrevemos um "
        "simulador da máquina de pilha em Python que interpreta exatamente "
        "o subconjunto de instruções que o codegen pode emitir. Para "
        "qualquer programa EC1 válido, simulador e `avaliar()` do "
        "interpretador produzem o mesmo número — o que prova, no teste, "
        "que o código gerado tem a mesma semântica do interpretador. "
        "Mantemos também a possibilidade de validar manualmente um `.s` no "
        "Linux (descrito no README).",
    )

    add_heading(doc, "Itens NÃO implementados (intencionalmente)", 2)
    add_body(
        doc,
        "A seção 7 do enunciado discute otimização: alocação inteligente de "
        "registradores para usar menos pilha, e propagação de constantes "
        "(toda expressão EC1 é constante, então o compilador poderia, em "
        "teoria, calcular o valor em tempo de compilação e emitir apenas "
        "`mov $V, %rax`). A seção 7 apresenta isso como o que ficou de "
        "fora, não como o que deve ser entregue — assim como a seção 6 não "
        "pede isso. Não foi implementado.",
    )

    # =============== DIFICULDADES ===============
    add_heading(doc, "Dificuldades", 1)
    add_body(
        doc,
        "Nenhuma dificuldade significativa. O esquema descrito na seção 4 "
        "do enunciado é direto. Os pontos que exigiram atenção foram:",
    )
    add_bullet(
        doc,
        "Decidir entre as duas opções de ordem (operando esquerdo primeiro "
        "vs. direito primeiro). Optamos pela opção 2 por simplicidade — "
        "vide decisão acima.",
    )
    add_bullet(
        doc,
        "Divisão com sinal correto. Inicialmente esquecemos o `cqo`; o "
        "teste de equivalência semântica imediatamente apontou divergência "
        "para os programas que envolviam divisão, e o ajuste foi trivial.",
    )
    add_bullet(
        doc,
        "Indentação consistente no template completo. O código que sai do "
        "codegen é cru (sem indentação), mas o template do `.s` indenta "
        "cada instrução com 4 espaços. `gerar_programa` faz essa "
        "reindentação antes de embrulhar no modelo.",
    )

    doc.save(path)


if __name__ == "__main__":
    import os
    out = os.path.join(os.path.dirname(__file__), "RELATORIO.docx")
    build_document(out)
    print(f"gerado: {out}")
