"""Gera o arquivo RELATORIO.docx a partir do conteudo do RELATORIO.md.

Uso:
    python gerar_relatorio.py

Requer: python-docx (pip install python-docx).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT = "Calibri"
BODY_SIZE = Pt(11)
CODE_FONT = "Consolas"


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.bold = True
    if level == 0:
        run.font.size = Pt(20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    else:
        run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12 if level > 0 else 0)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(6)
    add_runs_with_inline_code(p, text)
    return p


def add_runs_with_inline_code(p, text):
    """Quebra o texto por trechos entre crases (codigo inline) e adiciona como runs."""
    parts = text.split("`")
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.font.name = FONT
        run.font.size = BODY_SIZE
        if i % 2 == 1:
            run.font.name = CODE_FONT
            run.font.size = Pt(10)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_runs_with_inline_code(p, text)


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    add_runs_with_inline_code(p, text)


def add_centered_info_line(doc, label, value, bold_label=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    if label:
        r = p.add_run(label + " ")
        r.bold = bold_label
        r.font.name = FONT
        r.font.size = BODY_SIZE
    r2 = p.add_run(value)
    r2.font.name = FONT
    r2.font.size = BODY_SIZE


def build_document(path):
    doc = Document()

    # Estilo padrao
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = BODY_SIZE

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # Cabecalho institucional
    add_centered_info_line(doc, "", "Universidade Federal da Paraíba (UFPB)")
    add_centered_info_line(doc, "", "Centro de Informática – Curso de Ciência da Computação")
    add_centered_info_line(doc, "Disciplina:", "Construção de Compiladores 1")
    add_centered_info_line(doc, "Professor:", "Andrei de Araújo Formiga")
    doc.add_paragraph()

    # Titulo
    add_heading(doc, "Relatório – Atividade 02", 0)
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Compilador CI (Constantes Inteiras)")
    r_sub.italic = True
    r_sub.font.size = Pt(13)
    r_sub.font.name = FONT
    p_sub.paragraph_format.space_after = Pt(18)

    # Integrantes
    add_heading(doc, "Integrantes do grupo", 1)
    tabela = doc.add_table(rows=1, cols=2)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.style = "Light Grid Accent 1"
    hdr = tabela.rows[0].cells
    hdr[0].text = ""
    hdr[1].text = ""
    h1 = hdr[0].paragraphs[0].add_run("Nome")
    h2 = hdr[1].paragraphs[0].add_run("Matrícula")
    for run in (h1, h2):
        run.bold = True
        run.font.name = FONT
        run.font.size = BODY_SIZE
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for cell in hdr:
        set_cell_shading(cell, "1F3A5F")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    integrantes = [
        ("Davi Alves Rodrigues", "20230102377"),
        ("Gabriel Barbosa Ribeiro de Oliveira", "20230012814"),
        ("João Vitor Sampaio Costa", "20230089776"),
        ("Nathan Meira Nóbrega", "20240008904"),
    ]
    for nome, matricula in integrantes:
        row = tabela.add_row().cells
        for idx, valor in enumerate((nome, matricula)):
            row[idx].text = ""
            r = row[idx].paragraphs[0].add_run(valor)
            r.font.name = FONT
            r.font.size = BODY_SIZE
            row[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()

    # Como a atividade foi feita
    add_heading(doc, "Como a atividade foi feita", 1)
    add_body(
        doc,
        "O compilador foi implementado como um único script seguindo a estrutura "
        "clássica de análise → síntese descrita no enunciado:",
    )
    add_numbered(
        doc,
        "Leitura da entrada. O nome do arquivo `.ci` é recebido como argumento na "
        "linha de comando (`sys.argv[1]`) e o conteúdo é lido inteiro em memória.",
    )
    add_numbered(
        doc,
        "Análise. O conteúdo é validado contra a gramática da linguagem CI: após "
        "remover espaços em branco em volta, o texto precisa casar com a expressão "
        "regular `\\d+` (um ou mais dígitos). Qualquer outra coisa (letras, "
        "símbolos, string vazia) é tratada como erro de sintaxe — o compilador "
        "imprime uma mensagem no `stderr` e encerra com código de saída diferente "
        "de zero, sem gerar arquivo de saída.",
    )
    add_numbered(
        doc,
        "Síntese / geração de código. A constante validada é interpolada em um "
        "template de assembly x86-64 (sintaxe GNU Assembler) que segue exatamente "
        "o modelo do enunciado: seção `.text`, rótulo `_start`, a instrução "
        "`mov $<n>, %rax` gerada pelo compilador, chamadas a `imprime_num` e "
        "`sair`, e a inclusão final de `runtime.s`.",
    )
    add_numbered(
        doc,
        "Escrita da saída. O arquivo `.s` resultante é escrito no mesmo diretório "
        "da entrada, com o mesmo nome base e extensão trocada de `.ci` para `.s`.",
    )

    # Linguagem
    add_heading(doc, "Linguagem utilizada", 1)
    add_body(
        doc,
        "Python. A escolha foi por simplicidade: não há etapa de build, a leitura "
        "de arquivos e o uso de expressões regulares são diretos na biblioteca "
        "padrão, e o mesmo script roda em Windows e Linux sem alteração. Como a "
        "linguagem CI tem apenas dígitos, qualquer linguagem serviria; Python "
        "deixou o código do compilador curto e legível.",
    )

    # Dificuldades
    add_heading(doc, "Dificuldades", 1)
    add_body(
        doc,
        "Não houve dificuldades significativas. A linguagem CI é minimalista, "
        "então a análise se reduz a uma checagem de regex e a geração de código "
        "a uma interpolação de string. Os únicos pontos que exigiram alguma "
        "decisão foram:",
    )
    add_bullet(
        doc,
        "Definir o que é “espaço em branco aceitável” na entrada — optamos por "
        "permitir espaços/quebras de linha no início e no fim do arquivo (uso de "
        "`strip()`), o que é comum em editores de texto, mas exigir que o "
        "conteúdo efetivo seja apenas dígitos.",
    )
    add_bullet(
        doc,
        "Convenção do nome do arquivo de saída — escolhemos preservar o caminho "
        "da entrada e apenas trocar a extensão, o que torna a localização do "
        "`.s` previsível para quem usa o compilador.",
    )
    add_body(
        doc,
        "A verificação de overflow para 64 bits, mencionada como opcional no "
        "enunciado, não foi implementada.",
    )

    # Testes
    add_heading(doc, "Testes", 1)
    add_bullet(
        doc,
        "`testes/p1.ci` contém `42` — programa correto. O compilador gera "
        "`testes/p1.s` com `mov $42, %rax` no lugar indicado.",
    )
    add_bullet(
        doc,
        "`testes/erro.ci` contém `42abc` — programa com erro de sintaxe. O "
        "compilador encerra com código de saída 1 e imprime uma mensagem de "
        "erro no `stderr`, sem gerar arquivo de saída.",
    )

    doc.save(path)


if __name__ == "__main__":
    import os
    out = os.path.join(os.path.dirname(__file__), "RELATORIO.docx")
    build_document(out)
    print(f"gerado: {out}")
