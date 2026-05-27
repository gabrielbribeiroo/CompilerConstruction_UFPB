"""Gera o arquivo RELATORIO.docx a partir do conteudo do RELATORIO.md.

Uso:
    python gerar_relatorio.py

Requer: python-docx (pip install python-docx).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT = "Calibri"
BODY_SIZE = Pt(11)
CODE_FONT = "Consolas"
HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x5F)
TABLE_HEADER_BG = "1F3A5F"
CODE_BG = "F4F4F4"


# -------- helpers --------

def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_paragraph_shading(paragraph, hex_color):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.bold = True
    if level == 0:
        run.font.size = Pt(20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = HEADING_COLOR
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = HEADING_COLOR
    else:
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(14 if level > 0 else 0)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_runs_with_inline_code(p, text):
    """Quebra o texto por trechos entre crases e adiciona como runs separados."""
    parts = text.split("`")
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.font.name = FONT
        run.font.size = BODY_SIZE
        if i % 2 == 1:
            run.font.name = CODE_FONT
            run.font.size = Pt(10)


def add_body(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(6)
    add_runs_with_inline_code(p, text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_runs_with_inline_code(p, text)
    return p


def add_code_block(doc, code):
    """Bloco de codigo pre-formatado: monoespacada, fundo cinza claro."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    set_paragraph_shading(p, CODE_BG)
    linhas = code.rstrip("\n").split("\n")
    for i, linha in enumerate(linhas):
        if i > 0:
            run = p.add_run()
            run.add_break()
        run = p.add_run(linha)
        run.font.name = CODE_FONT
        run.font.size = Pt(9.5)
    return p


def add_centered_info_line(doc, label, value, bold_label=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    if label:
        r = p.add_run(label + " ")
        r.bold = bold_label
        r.font.name = FONT
        r.font.size = BODY_SIZE
    r2 = p.add_run(value)
    r2.font.name = FONT
    r2.font.size = BODY_SIZE


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for idx, h in enumerate(headers):
        hdr_cells[idx].text = ""
        r = hdr_cells[idx].paragraphs[0].add_run(h)
        r.bold = True
        r.font.name = FONT
        r.font.size = BODY_SIZE
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[idx], TABLE_HEADER_BG)
        hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if col_widths:
            hdr_cells[idx].width = col_widths[idx]
    for row in rows:
        cells = table.add_row().cells
        for idx, valor in enumerate(row):
            cells[idx].text = ""
            add_runs_with_inline_code(cells[idx].paragraphs[0], str(valor))
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if col_widths:
                cells[idx].width = col_widths[idx]
    return table


# -------- documento --------

def build_document(path):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = BODY_SIZE

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # Cabecalho institucional
    add_centered_info_line(doc, "", "Universidade Federal da Paraíba (UFPB)")
    add_centered_info_line(doc, "", "Centro de Informática – Curso de Ciência da Computação")
    add_centered_info_line(doc, "Disciplina:", "Construção de Compiladores 1")
    add_centered_info_line(doc, "Professor:", "Andrei de Araújo Formiga")
    doc.add_paragraph()

    # Titulo
    add_heading(doc, "Relatório – Atividade 05", 0)
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Análise Sintática EC1 (Expressões Constantes 1)")
    r_sub.italic = True
    r_sub.font.size = Pt(13)
    r_sub.font.name = FONT
    p_sub.paragraph_format.space_after = Pt(18)

    # Integrantes
    add_heading(doc, "Integrantes do grupo", 1)
    add_table(
        doc,
        headers=["Nome", "Matrícula"],
        rows=[
            ["Davi Alves Rodrigues", "20230102377"],
            ["Gabriel Barbosa Ribeiro de Oliveira", "20230012814"],
            ["João Vitor Sampaio Costa", "20230089776"],
            ["Nathan Meira Nóbrega", "20240008904"],
        ],
    )
    doc.add_paragraph()

    # =============== O QUE FOI IMPLEMENTADO ===============
    add_heading(doc, "O que foi implementado", 1)

    # 1. AST
    add_heading(doc, "1. Árvore de sintaxe abstrata (ast_ec1.py)", 2)
    add_body(
        doc,
        "Seguindo a sugestão da seção 5 do enunciado, a árvore é representada "
        "por uma hierarquia de classes:",
    )
    add_bullet(doc, "`Exp` — classe-base abstrata. Define a interface `avaliar()` e `__str__()`.")
    add_bullet(doc, "`Const(valor: int)` — folha que carrega o valor de uma constante inteira.")
    add_bullet(
        doc,
        "`OpBin(op: Op, esq: Exp, dir: Exp)` — nó interno com operador e dois "
        "filhos (operando esquerdo e direito).",
    )
    add_bullet(
        doc,
        "`Op` — enumeração com os quatro operadores (`SOMA`, `SUB`, `MULT`, "
        "`DIV`), com `value` igual ao símbolo (`+`, `-`, `*`, `/`) para "
        "facilitar a reconstrução textual.",
    )
    add_body(
        doc,
        "`Const` e `OpBin` são `dataclasses` com `frozen=True`. Isso fornece "
        "automaticamente `__eq__` por valor, o que permite comparações diretas "
        "entre ASTs nos testes (`self.assertEqual(arvore, esperado)`) sem "
        "código extra, e também torna os nós hasháveis.",
    )

    # 2. Parser
    add_heading(doc, "2. Analisador sintático descendente recursivo (parser.py)", 2)
    add_body(
        doc,
        "A classe `AnalisadorSintatico` implementa o algoritmo descrito nas "
        "seções 3 e 6 do enunciado. Para cada não-terminal há uma função "
        "privada:",
    )
    add_bullet(
        doc,
        "`_analisa_exp()` — implementa a regra "
        "`<expressao> ::= <literal> | '(' <expressao> <operador> <expressao> ')'`. "
        "Consome o próximo token e ramifica conforme o tipo: `NUMERO` vira "
        "`Const`; `PAREN_ESQ` dispara a análise da operação binária (esquerda, "
        "operador, direita, fechamento); qualquer outra coisa é erro.",
    )
    add_bullet(
        doc,
        "`_analisa_operador()` — implementa "
        "`<operador> ::= '+' | '-' | '*' | '/'`, mapeando o tipo do token "
        "para o `Op` correspondente.",
    )
    add_bullet(
        doc,
        "`analisa_programa()` — chama `_analisa_exp()` e, ao final, exige que "
        "o próximo token seja `EOF`, garantindo que não há lixo após a expressão.",
    )
    add_body(
        doc,
        "Erros sintáticos são reportados via exceção `ErroSintatico`, que "
        "carrega a posição (índice do caractere no texto-fonte, herdada do "
        "token problemático) e uma mensagem descrevendo o que era esperado "
        "versus o que foi encontrado.",
    )

    # 3. Interpretador
    add_heading(doc, "3. Interpretador por varredura da árvore (avaliar())", 2)
    add_body(
        doc,
        "Conforme sugerido na seção 8 do enunciado, o interpretador é "
        "implementado como o método `avaliar()` na classe base `Exp`, "
        "sobrescrito nas subclasses:",
    )
    add_bullet(doc, "`Const.avaliar()` devolve o valor literal.")
    add_bullet(
        doc,
        "`OpBin.avaliar()` avalia recursivamente os dois filhos e aplica o "
        "operador.",
    )
    add_body(
        doc,
        "A divisão usa truncamento para zero (`int(a / b)`), coerente com a "
        "semântica de divisão inteira de assembly — alvo das próximas "
        "atividades. Divisão por zero levanta `ZeroDivisionError`.",
    )

    # 4. Lexer
    add_heading(doc, "4. Analisador léxico (lexer.py)", 2)
    add_body(
        doc,
        "Reaproveita a estrutura da Atividade 04: `TipoToken` enum, `Token` "
        "como `dataclass`, exceção `ErroLexico` com posição e caractere, e a "
        "classe `AnalisadorLexico` com as duas interfaces discutidas na seção "
        "7 do enunciado:",
    )
    add_bullet(doc, "`proximo_token()` — consome e retorna o próximo token.")
    add_bullet(
        doc,
        "`olhar_proximo()` — peek sem consumir (usado nos testes; o parser "
        "funciona bem só com `proximo_token`).",
    )
    add_body(
        doc,
        "Espaços, tabulações, quebras de linha e CR são ignorados; "
        "comentários de linha (`# ...`) também (extensão herdada da atividade "
        "anterior).",
    )

    # 5. Ponto de entrada
    add_heading(doc, "5. Ponto de entrada (ec1.py)", 2)
    add_body(
        doc,
        "Recebe o nome do arquivo `.ec1` na linha de comando, executa "
        "lex → parse → avalia e imprime o valor em `stdout`. Erros de E/S, "
        "léxicos, sintáticos e de divisão por zero vão para `stderr` com "
        "código de saída 1.",
    )

    # 6. Suite de testes
    add_heading(doc, "6. Suíte de testes (tests/test_parser.py)", 2)
    add_body(doc, "34 testes, 0 falhas, divididos em quatro classes do módulo `unittest`:")
    add_table(
        doc,
        headers=["Classe", "Foco", "Casos"],
        rows=[
            [
                "TestArvoreSintatica",
                "Estrutura da AST produzida pelo parser",
                "11",
            ],
            [
                "TestInterpretador",
                "Valor numérico produzido por avaliar()",
                "9",
            ],
            [
                "TestErrosSintaticos",
                "Detecção e reporte correto de programas mal formados",
                "10",
            ],
            [
                "TestImpressaoCanonica",
                "Reconstituição da expressão via __str__()",
                "4",
            ],
            ["Total", "", "34"],
        ],
    )
    doc.add_paragraph()
    add_body(doc, "Destaques:")
    add_bullet(
        doc,
        "O exemplo literal do enunciado (`(33 + (912 * 11))`) tem dois testes "
        "dedicados: um verifica que a AST construída é exatamente "
        "`OpBin(SOMA, Const(33), OpBin(MULT, Const(912), Const(11)))` e outro "
        "verifica que `avaliar()` devolve `10065`.",
    )
    add_bullet(
        doc,
        "A expressão complexa `((427 / 7) + (11 * (231 + 5)))` também é "
        "verificada estrutural e numericamente (resultado `2657`).",
    )
    add_bullet(
        doc,
        "Os testes de erro verificam que o `ErroSintatico` é levantado para "
        "parêntese não fechado, parêntese fechando inesperado, operador "
        "faltando ou substituído, operando faltando, entrada vazia, entrada "
        "só com espaços e lixo após a expressão raiz.",
    )

    # 7. Arquivos de exemplo
    add_heading(doc, "7. Arquivos de exemplo", 2)
    add_body(
        doc,
        "Quatro programas válidos (`valido1.ec1` a `valido4.ec1`, cobrindo "
        "constante simples, multiplicação simples e os dois exemplos do "
        "enunciado) e três com erro sintático em pontos diferentes da "
        "árvore (`invalido_parenteses.ec1`, `invalido_operador.ec1`, "
        "`invalido_lixo_no_fim.ec1`).",
    )
    add_table(
        doc,
        headers=["Arquivo", "Conteúdo", "Resultado esperado"],
        rows=[
            ["valido1.ec1", "333", "333"],
            ["valido2.ec1", "(6 * 7)", "42"],
            ["valido3.ec1", "(33 + (912 * 11))", "10065"],
            [
                "valido4.ec1",
                "((427 / 7) + (11 * (231 + 5)))",
                "2657",
            ],
            [
                "invalido_parenteses.ec1",
                "(33 + (912 * 11)",
                "erro sintático, exit 1",
            ],
            [
                "invalido_operador.ec1",
                "(3 3 4)",
                "erro sintático, exit 1",
            ],
            [
                "invalido_lixo_no_fim.ec1",
                "(6 * 7) 42",
                "erro sintático, exit 1",
            ],
        ],
    )

    # =============== ESTRUTURA DE ARQUIVOS ===============
    add_heading(doc, "Estrutura de arquivos entregue", 1)
    add_code_block(
        doc,
        "analise-sintatica-ec1/\n"
        "├── lexer.py\n"
        "├── ast_ec1.py\n"
        "├── parser.py\n"
        "├── ec1.py\n"
        "├── tests/\n"
        "│   └── test_parser.py\n"
        "├── exemplos/\n"
        "│   ├── valido1.ec1\n"
        "│   ├── valido2.ec1\n"
        "│   ├── valido3.ec1\n"
        "│   ├── valido4.ec1\n"
        "│   ├── invalido_parenteses.ec1\n"
        "│   ├── invalido_operador.ec1\n"
        "│   └── invalido_lixo_no_fim.ec1\n"
        "├── README.md\n"
        "├── PLANO.md\n"
        "└── RELATORIO.md",
    )

    # =============== DECISOES DE PROJETO ===============
    add_heading(doc, "Decisões de projeto", 1)

    add_heading(doc, "Por que Const e OpBin como dataclass(frozen=True)?", 2)
    add_body(
        doc,
        "Por dois motivos: o `__eq__` automático por valor encurta "
        "drasticamente os testes de estrutura da AST, e `frozen=True` deixa "
        "explícito que os nós são imutáveis — a AST construída pelo parser "
        "não é modificada por ninguém depois disso, o que é um invariante "
        "útil para o interpretador e qualquer estágio futuro do compilador.",
    )

    add_heading(doc, "Por que Op como Enum com value igual ao símbolo?", 2)
    add_body(
        doc,
        "Um único valor representa “qual operador” — não há razão para "
        "misturar isso com o tipo de token. Manter o símbolo como `value` "
        "deixa `OpBin.__str__` trivial "
        "(`f\"({self.esq} {self.op.value} {self.dir})\"`) e ajuda a debugar.",
    )

    add_heading(
        doc,
        "Por que avaliar() como método nas classes da AST, em vez de uma função externa?",
        2,
    )
    add_body(
        doc,
        "É a abordagem sugerida pelo enunciado (seção 8) e dispensa "
        "`isinstance` ou `match` em uma função à parte. Adicionar um novo nó "
        "no futuro (por exemplo, operação unária) exige apenas implementar "
        "`avaliar()` na nova classe.",
    )

    add_heading(doc, "Por que o parser usa só proximo_token() (consume), sem peek?", 2)
    add_body(
        doc,
        "O pseudo-código do enunciado consome o próximo token e ramifica "
        "pelo tipo, sem precisar devolvê-lo. Como o token de abertura "
        "(`PAREN_ESQ`) e o literal não precisam ser preservados após a "
        "decisão (são “gastos” na própria decisão), isso simplifica o "
        "código e não custa nada.",
    )

    add_heading(doc, "Verificação de fim da entrada", 2)
    add_body(
        doc,
        "`analisa_programa()` exige `EOF` após a expressão raiz para detectar "
        "programas como `(6 * 7) 42` — entrada perfeitamente válida do ponto "
        "de vista de `_analisa_exp()` para o trecho `(6 * 7)`, mas que tem "
        "lixo depois. Sem essa checagem, o parser silenciaria erros desse "
        "tipo.",
    )

    add_heading(doc, "Itens NÃO implementados (intencionalmente)", 2)
    add_body(
        doc,
        "O enunciado menciona na seção 9 a possibilidade de imprimir a "
        "árvore em formato ASCII visual ou via graphviz — explicitamente "
        "como “outras possibilidades”. Esses itens não fazem parte do "
        "conjunto pedido na seção 10 (artefato para entrega) e não foram "
        "implementados. A impressão textual simples (reconstrução da "
        "expressão entre parênteses) é implementada via `__str__()` nas "
        "classes da AST porque é útil para testes e tem custo desprezível.",
    )

    # =============== DIFICULDADES ===============
    add_heading(doc, "Dificuldades", 1)
    add_body(
        doc,
        "Nenhuma dificuldade significativa. O pseudocódigo da seção 6 do "
        "enunciado descreve o parser de forma quase completa; a tradução "
        "para Python foi direta. Os únicos pontos que exigiram atenção:",
    )
    add_bullet(
        doc,
        "Decidir onde verificar o fim da entrada — a função `_analisa_exp` é "
        "recursiva e é chamada também a partir de si mesma, então fazer a "
        "verificação dentro dela quebraria os casos aninhados. A solução foi "
        "isolar essa verificação em `analisa_programa()`, que envolve o "
        "`_analisa_exp` raiz.",
    )
    add_bullet(
        doc,
        "Tipagem da divisão — como a linguagem só tem inteiros e as próximas "
        "etapas alvejam assembly inteiro, optamos por divisão inteira com "
        "truncamento para zero (`int(a / b)`) em vez de divisão real "
        "(`a / b`) ou divisão floor (`a // b`, que arredonda para -inf). O "
        "truncamento para zero é o que `idiv` faz em x86-64.",
    )

    doc.save(path)


if __name__ == "__main__":
    import os
    out = os.path.join(os.path.dirname(__file__), "RELATORIO.docx")
    build_document(out)
    print(f"gerado: {out}")
