"""Gera o arquivo RELATORIO.docx a partir do conteudo do RELATORIO.md.

Uso:
    python gerar_relatorio.py

Requer: python-docx (pip install python-docx).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT = "Calibri"
BODY_SIZE = Pt(11)
CODE_FONT = "Consolas"
HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x5F)
TABLE_HEADER_BG = "1F3A5F"
CODE_BG = "F4F4F4"


# -------- helpers --------

def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_paragraph_shading(paragraph, hex_color):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.bold = True
    if level == 0:
        run.font.size = Pt(20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = HEADING_COLOR
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = HEADING_COLOR
    else:
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(14 if level > 0 else 0)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_runs_with_inline_code(p, text):
    parts = text.split("`")
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.font.name = FONT
        run.font.size = BODY_SIZE
        if i % 2 == 1:
            run.font.name = CODE_FONT
            run.font.size = Pt(10)


def add_body(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(6)
    add_runs_with_inline_code(p, text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_runs_with_inline_code(p, text)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    set_paragraph_shading(p, CODE_BG)
    linhas = code.rstrip("\n").split("\n")
    for i, linha in enumerate(linhas):
        if i > 0:
            run = p.add_run()
            run.add_break()
        run = p.add_run(linha)
        run.font.name = CODE_FONT
        run.font.size = Pt(9.5)
    return p


def add_centered_info_line(doc, label, value, bold_label=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    if label:
        r = p.add_run(label + " ")
        r.bold = bold_label
        r.font.name = FONT
        r.font.size = BODY_SIZE
    r2 = p.add_run(value)
    r2.font.name = FONT
    r2.font.size = BODY_SIZE


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for idx, h in enumerate(headers):
        hdr_cells[idx].text = ""
        r = hdr_cells[idx].paragraphs[0].add_run(h)
        r.bold = True
        r.font.name = FONT
        r.font.size = BODY_SIZE
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[idx], TABLE_HEADER_BG)
        hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if col_widths:
            hdr_cells[idx].width = col_widths[idx]
    for row in rows:
        cells = table.add_row().cells
        for idx, valor in enumerate(row):
            cells[idx].text = ""
            add_runs_with_inline_code(cells[idx].paragraphs[0], str(valor))
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if col_widths:
                cells[idx].width = col_widths[idx]
    return table


# -------- documento --------

def build_document(path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = BODY_SIZE
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # Cabecalho institucional
    add_centered_info_line(doc, "", "Universidade Federal da Paraíba (UFPB)")
    add_centered_info_line(doc, "", "Centro de Informática – Curso de Ciência da Computação")
    add_centered_info_line(doc, "Disciplina:", "Construção de Compiladores 1")
    add_centered_info_line(doc, "Professor:", "Andrei de Araújo Formiga")
    doc.add_paragraph()

    # Titulo
    add_heading(doc, "Relatório – Atividade 10", 0)
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Compilador Fun — Funções")
    r_sub.italic = True
    r_sub.font.size = Pt(13)
    r_sub.font.name = FONT
    p_sub.paragraph_format.space_after = Pt(18)

    # Integrantes
    add_heading(doc, "Integrantes do grupo", 1)
    add_table(
        doc,
        headers=["Nome", "Matrícula"],
        rows=[
            ["Davi Alves Rodrigues", "20230102377"],
            ["Gabriel Barbosa Ribeiro de Oliveira", "20230012814"],
            ["João Vitor Sampaio Costa", "20230089776"],
            ["Nathan Meira Nóbrega", "20240008904"],
        ],
    )
    doc.add_paragraph()

    # =============== O QUE FOI IMPLEMENTADO ===============
    add_heading(doc, "O que foi implementado", 1)

    # 1. A linguagem Fun
    add_heading(doc, "1. A linguagem Fun: funções com pilha de ativação", 2)
    add_body(
        doc,
        "Cmd (Atividade 09) já era Turing-completa, mas não tinha "
        "nenhuma forma de abstração/reuso de código: todo cálculo "
        "tinha que ser escrito por extenso no corpo do programa. Fun "
        "estende Cmd com declaração e chamada de funções — com "
        "parâmetros, variáveis locais próprias e recursão direta — a "
        "primeira mudança neste conjunto de compiladores que exige "
        "gerar um registro de ativação de verdade na pilha do sistema, "
        "em vez de só usá-la como área de rascunho para avaliar "
        "expressões.",
    )
    add_body(
        doc,
        "Um programa Fun é `<decl>* 'main' '{' <cmd>* 'return' <exp> "
        "';' '}'`: zero ou mais declarações — de variável global "
        "(`var x = exp;`) ou de função (`fun nome(params) { ... }`) — "
        "seguidas do bloco `main`. Dentro de uma função, `var` declara "
        "uma variável local, visível só naquela função e que esconde "
        "uma global de mesmo nome, se houver.",
    )

    # 2. Alteracoes lexicas
    add_heading(doc, "2. Alterações léxicas (lexer.py)", 2)
    add_body(
        doc,
        "Um token novo, `VIRGULA` (`,`), para separar parâmetros "
        "formais e argumentos de chamada, e três palavras-chave "
        "novas: `fun`, `var`, `main`. Reconhecidas pela mesma técnica "
        "das atividades anteriores (monta o identificador inteiro "
        "primeiro, só depois compara contra `PALAVRAS_CHAVE`), "
        "preservando o comportamento de `ifx`-como-identificador já "
        "testado desde a Atividade 09.",
    )

    # 3. AST
    add_heading(doc, "3. Árvore de sintaxe abstrata (ast_fun.py)", 2)
    add_body(
        doc,
        "Reaproveita `Exp`/`Const`/`Var`/`OpBin`/`Op` e "
        "`Cmd`/`Atrib`/`If`/`While` de Cmd sem alteração de forma. "
        "Nós novos:",
    )
    add_bullet(
        doc,
        "`Chamada(nome, args, posicao)` — uma expressão, no mesmo "
        "nível de `Var`/`Const`/`OpBin`. `avaliar()` é o coração do "
        "interpretador de referência para funções: busca a `FunDecl` "
        "no `Contexto`, avalia cada argumento no ambiente do "
        "chamador, cria um `Ambiente` novo com os parâmetros já "
        "ligados aos valores avaliados, inicializa as variáveis "
        "locais da função nesse ambiente novo (sequencialmente, cada "
        "uma pode referenciar as anteriores), executa os comandos do "
        "corpo e devolve o valor da expressão final — tudo isolado do "
        "ambiente de quem chamou.",
    )
    add_bullet(
        doc,
        "`VarDecl(nome, exp)` — usado tanto para variáveis globais "
        "(no nível do programa) quanto locais (no início do corpo de "
        "uma função).",
    )
    add_bullet(doc, "`FunDecl(nome, params, vardecls, comandos, exp_final)`.")
    add_bullet(
        doc,
        "`Programa(declaracoes, comandos, exp_final)` — `declaracoes` "
        "agora é uma sequência heterogênea de `VarDecl`/`FunDecl`.",
    )
    add_body(
        doc,
        "Dois objetos novos dão suporte à avaliação: `Contexto` "
        "(globais + funções declaradas, compartilhado por todas as "
        "chamadas de um mesmo programa) e `Ambiente` (locais de uma "
        "chamada específica + o `Contexto`; `obter`/`definir` "
        "primeiro tentam o dicionário local, só recaindo para o "
        "global se o nome não for um parâmetro/local daquela chamada "
        "— implementando o shadowing).",
    )

    # 4. Sintaxe
    add_heading(doc, "4. Alterações na sintaxe (parser.py)", 2)
    add_body(
        doc,
        "`analisa_programa()` monta a lista de declarações espiando "
        "(peek) o próximo token: `var` inicia uma `VarDecl` de nível "
        "de programa, `fun` inicia uma `FunDecl`; qualquer outro "
        "token encerra a lista de declarações e exige `main`.",
    )
    add_body(
        doc,
        "A mudança mais delicada da atividade é distinguir uma "
        "referência a variável de uma chamada de função dentro de "
        "`_analisa_prim()`: os dois começam com um `IDENT`. A solução "
        "é um lookahead de exatamente um token depois do "
        "identificador consumido — se for `(`, é uma `Chamada` (e o "
        "parser entra em `_analisa_lista_parametros_reais()`, que "
        "reaproveita o mesmo padrão de laço com vírgula já usado para "
        "os parâmetros formais); caso contrário, é um `Var`. Como o "
        "lexer já expõe `olhar_proximo()` (peek sem consumir) desde a "
        "Atividade 06, essa decisão não exige nenhum retrocesso "
        "(backtracking) no parser.",
    )

    # 5. Semantica
    add_heading(doc, "5. Análise semântica (semantica.py)", 2)
    add_body(
        doc,
        "`TabelaSimbolos` ganha dois espaços de nomes: "
        "`_vars_globais` (um set) e `_funcoes` (um dict nome → "
        "aridade). A verificação de uma `FunDecl` "
        "(`_verifica_fundecl`) constrói um conjunto de nomes locais "
        "começando pelos parâmetros, processa as vardecls do corpo "
        "sequencialmente (cada uma é verificada contra os nomes "
        "locais já vistos até ali, depois some ao conjunto — "
        "impedindo que uma variável local se referencie antes de ser "
        "declarada), e por fim verifica comandos e expressão final "
        "contra esse escopo local (caindo para o escopo global quando "
        "o nome não é local, implementando o mesmo shadowing do "
        "interpretador).",
    )
    add_body(
        doc,
        "O detalhe que permite recursão direta, mas não mútua, é a "
        "ordem de duas linhas em `verifica_programa`:",
    )
    add_code_block(
        doc,
        "elif isinstance(decl, FunDecl):\n"
        "    # registra a funcao ANTES de verificar seu proprio corpo,\n"
        "    # para permitir chamadas diretamente recursivas\n"
        "    tabela.declarar_funcao(decl.nome, len(decl.params))\n"
        "    _verifica_fundecl(decl, tabela)",
    )
    add_body(
        doc,
        "Como o nome da função já está na tabela quando seu próprio "
        "corpo é verificado, uma chamada recursiva a si mesma é "
        "aceita. Como as declarações são processadas sequencialmente, "
        "uma função `a` que chama `b`, declarada depois dela no "
        "texto, é rejeitada (`b` ainda não está na tabela) — "
        "exatamente o comportamento que o enunciado prevê "
        "implicitamente ao tratar recursão mútua como extensão "
        "opcional (seção 5.1), e que confirmamos com um teste "
        "dedicado (`test_funcao_mutuamente_recursiva_e_rejeitada`).",
    )

    # 6. Codegen
    add_heading(doc, "6. Geração de código (codegen.py)", 2)
    add_heading(doc, "Chamada de função (seção 6.1.1–6.1.2 do enunciado)", 2)
    add_code_block(
        doc,
        "<codigo do ultimo argumento>; push %rax\n"
        "...\n"
        "<codigo do primeiro argumento>; push %rax\n"
        "call <nome>\n"
        "add $8*N, %rsp     # pulado se N == 0 (limpeza feita pelo CHAMADOR)",
    )
    add_body(
        doc,
        "Os argumentos são empilhados em ordem inversa (último "
        "primeiro), de modo que, dentro da função, o primeiro "
        "parâmetro fique mais perto do topo da pilha na hora de "
        "calcular seu deslocamento — o esquema descrito na seção "
        "6.1.3 do enunciado.",
    )
    add_heading(doc, "Prólogo e epílogo da função (seção 6.1.4 do enunciado)", 2)
    add_code_block(
        doc,
        "<nome>:\n"
        "    push %rbp\n"
        "    sub $8*L, %rsp      # pulado se L == 0\n"
        "    mov %rsp, %rbp\n"
        "    <codigo de cada variavel local, grava em deslocamento(%rbp)>\n"
        "    <codigo dos comandos>\n"
        "    <codigo da expressao final, deixa o valor em %rax>\n"
        "    add $8*L, %rsp      # pulado se L == 0\n"
        "    pop %rbp\n"
        "    ret",
    )
    add_body(
        doc,
        "Seguimos literalmente as duas instruções de epílogo do "
        "enunciado (`add $8*L, %rsp` seguido de `pop %rbp`) em vez da "
        "instrução `leave` do x86-64 — ver “Decisões de projeto” "
        "abaixo para o motivo, que só foi descoberto durante os "
        "testes de equivalência semântica.",
    )
    add_heading(doc, "Deslocamentos (seção 6.1.3 do enunciado)", 2)
    add_body(
        doc,
        "Uma variável local de índice `j` (0-based, ordem de "
        "declaração) fica em `8*j(%rbp)`; um parâmetro de índice `i` "
        "fica em `(8*L + 16 + 8*i)(%rbp)`, onde `L` é o número de "
        "locais da função — os `16` bytes cobrem o `%rbp` antigo e o "
        "endereço de retorno, ambos empilhados entre o fim dos "
        "argumentos e o início do corpo da função. `_calcula_quadro()` "
        "monta esse mapa nome → deslocamento uma vez por função, "
        "antes de gerar seu corpo; `_emit_carrega_variavel`/"
        "`_emit_armazena_variavel` consultam esse mapa primeiro e só "
        "caem para o símbolo `.bss` global (como em Cmd) quando o "
        "nome não é um local/parâmetro da função sendo gerada no "
        "momento (ou quando o gerador está processando o bloco "
        "`main`, onde o mapa é sempre `None`).",
    )
    add_heading(doc, "Inicialização de variáveis globais", 2)
    add_body(
        doc,
        "Uma variável global (`var x = exp;` no nível do programa) "
        "precisa que sua expressão inicial seja calculada em tempo "
        "de execução e armazenada no símbolo `.bss` correspondente "
        "antes que qualquer parte do programa a leia — `.lcomm` só "
        "reserva memória zerada, não executa código nenhum. O bloco "
        "principal (`_start`) agora começa emitindo o código de cada "
        "`VarDecl` de nível de programa, na ordem declarada, antes do "
        "código dos comandos de `main` — ver “Decisões de projeto” "
        "para o relato de como essa omissão foi descoberta.",
    )
    add_heading(doc, "Modelo de saída (seção 7 do enunciado)", 2)
    add_body(
        doc,
        "O .s gerado agora tem três blocos, nesta ordem: seção `.bss` "
        "(variáveis globais), o bloco principal (rótulo `_start`, "
        "terminando em `call imprime_num`/`call sair`), e o código de "
        "cada função declarada (rótulo com o nome da função, prólogo, "
        "corpo, epílogo) — antes do `.include \"runtime.s\"`, "
        "reaproveitado sem alteração desde a Atividade 02.",
    )

    # 7. CLI
    add_heading(doc, "7. Ponto de entrada (compfun.py)", 2)
    add_body(
        doc,
        "Mesma estrutura de sempre: lê o arquivo `.fun` → "
        "`analisar()` (léxico + sintático) → `verifica_programa()` "
        "(semântico) → `gerar_programa()` (codegen) → grava `.s`. "
        "Qualquer uma das três exceções (`ErroLexico`, "
        "`ErroSintatico`, `ErroSemantico`) é capturada, a mensagem "
        "vai para `stderr`, e o processo encerra com exit code 1 sem "
        "gravar `.s`.",
    )

    # 8. Variacao sintatica
    add_heading(doc, "8. Variação sintática", 2)
    add_body(
        doc,
        "Seguimos exatamente a sintaxe descrita no enunciado, sem "
        "implementar funções mutuamente recursivas nem funções locais "
        "(aninhadas dentro de outra função) — ambas citadas como "
        "extensões possíveis, não como parte do artefato mínimo "
        "pedido. Nenhuma variação sintática das atividades anteriores "
        "(operadores `<=`/`>=`, booleanos, etc.) foi revisitada.",
    )

    # 9. Testes
    add_heading(doc, "9. Suíte de testes (tests/test_fun.py)", 2)
    add_body(doc, "40 testes, 0 falhas, em 8 classes:")
    add_table(
        doc,
        headers=["Classe", "Foco", "Casos"],
        rows=[
            ["TestLexico", "Token de vírgula, palavras-chave novas", "3"],
            [
                "TestParser",
                "FunDecl/VarDecl, Chamada distinguida de Var, listas de parâmetros/argumentos",
                "8",
            ],
            [
                "TestErrosSintaticos",
                "Entradas mal formadas (parênteses não fechados, main ausente, etc.)",
                "5",
            ],
            [
                "TestSemantica",
                "Função/variável não declarada, aridade incorreta, escopo local, recursão direta permitida e mútua rejeitada",
                "8",
            ],
            [
                "TestInterpretacao",
                "Programa.avaliar() para abs, fib, chamada entre funções, variável local escondendo global",
                "5",
            ],
            [
                "TestEquivalenciaSemantica",
                "Código gerado bate com o interpretador (8 programas, incluindo recursão)",
                "1",
            ],
            [
                "TestCodegen",
                "Prólogo/epílogo, deslocamentos, limpeza de pilha pelo chamador",
                "6",
            ],
            [
                "TestCLI",
                "Comportamento de compfun.py como subprocesso, inclusive função recursiva",
                "4",
            ],
            ["Total", "", "40"],
        ],
    )
    doc.add_paragraph()
    add_body(doc, "Destaques:")
    add_bullet(
        doc,
        "`TestEquivalenciaSemantica` é, de longe, o teste mais "
        "importante desta atividade — e o que efetivamente pegou dois "
        "bugs reais antes da entrega (ver “Decisões de projeto”). O "
        "simulador (`simular_programa`) modela a pilha do sistema como "
        "memória endereçável de verdade (um dict endereço → valor, "
        "com `%rsp`/`%rbp` como inteiros reais que crescem/decrescem "
        "exatamente como no x86-64), em vez de reimplementar a "
        "aritmética de deslocamentos “por fora” com índices de lista "
        "— isso reduz o simulador a uma tradução quase literal de "
        "cada instrução, e foi o que tornou possível notar, instrução "
        "por instrução, exatamente onde a pilha ficava desbalanceada. "
        "Suporta `call`/`ret` com pilha de retorno própria, o que "
        "permite testar recursão de verdade (`fib(10)` chamando a si "
        "mesma).",
    )
    add_bullet(
        doc,
        "`test_recursao_direta_permitida` e "
        "`test_funcao_mutuamente_recursiva_e_rejeitada` cobrem "
        "exatamente a fronteira de escopo desta atividade (recursão "
        "direta dentro, recursão mútua fora).",
    )
    add_bullet(
        doc,
        "`test_limpeza_da_pilha_apos_chamada` e "
        "`test_sem_parametros_nao_gera_limpeza_de_pilha` confirmam o "
        "`add $8*N, %rsp` do lado do chamador, incluindo o caso "
        "`N == 0` (não deve gerar `add $0, %rsp` nem nenhuma instrução "
        "de limpeza).",
    )

    # 10. Exemplos
    add_heading(doc, "10. Arquivos de exemplo", 2)
    add_body(
        doc,
        "Quatro programas válidos e três inválidos. Os válidos: "
        "abs() e fib() recursivo (ambos exemplos literais do "
        "enunciado), uma função chamando outra (somaDosQuadrados "
        "chamando quadrado duas vezes), e um parâmetro de função "
        "escondendo uma variável global de mesmo nome (dobro(x) com "
        "um x global também declarado). Os inválidos: chamada a "
        "função nunca declarada, chamada com número de argumentos "
        "diferente da declaração, e uso em main de uma variável que "
        "só existe como local de outra função.",
    )

    # =============== EXEMPLO DE SAIDA ===============
    add_heading(doc, "Exemplo de saída — fib recursivo (trecho)", 1)
    add_code_block(
        doc,
        "_start:\n"
        "    # return fib(10);\n"
        "    mov $10, %rax\n"
        "    push %rax\n"
        "    call fib\n"
        "    add $8, %rsp\n"
        "    call imprime_num\n"
        "    call sair\n"
        "\n"
        "fib:\n"
        "    push %rbp\n"
        "    sub $8, %rsp\n"
        "    mov %rsp, %rbp\n"
        "    # var res = 0;\n"
        "    mov $0, %rax\n"
        "    mov %rax, 0(%rbp)\n"
        "    # if (n < 2) {\n"
        "    ...\n"
        "    jz Lfalso0\n"
        "    # res = 1;\n"
        "    mov $1, %rax\n"
        "    mov %rax, 0(%rbp)\n"
        "    jmp Lfim0\n"
        "Lfalso0:\n"
        "    # } else {\n"
        "    # res = (fib((n - 1)) + fib((n - 2)));\n"
        "    ...\n"
        "    call fib\n"
        "    add $8, %rsp\n"
        "    push %rax\n"
        "    ...\n"
        "    call fib\n"
        "    add $8, %rsp\n"
        "    pop %rbx\n"
        "    add %rbx, %rax\n"
        "    mov %rax, 0(%rbp)\n"
        "Lfim0:\n"
        "    # }\n"
        "    # return res;\n"
        "    mov 0(%rbp), %rax\n"
        "    add $8, %rsp\n"
        "    pop %rbp\n"
        "    ret",
    )
    add_body(
        doc,
        "Executado, imprime `89` (fib(0) = fib(1) = 1, então "
        "fib(10) = 89 nessa numeração).",
    )

    # =============== ESTRUTURA DE ARQUIVOS ===============
    add_heading(doc, "Estrutura de arquivos entregue", 1)
    add_code_block(
        doc,
        "compilador-fun/\n"
        "├── lexer.py\n"
        "├── ast_fun.py\n"
        "├── parser.py\n"
        "├── semantica.py\n"
        "├── codegen.py\n"
        "├── compfun.py\n"
        "├── runtime.s\n"
        "├── exemplos/\n"
        "│   ├── valido1.fun\n"
        "│   ├── valido2.fun\n"
        "│   ├── valido3.fun\n"
        "│   ├── valido4.fun\n"
        "│   ├── invalido_funcao_nao_declarada.fun\n"
        "│   ├── invalido_numero_de_parametros.fun\n"
        "│   └── invalido_variavel_fora_de_escopo.fun\n"
        "├── tests/\n"
        "│   └── test_fun.py\n"
        "├── README.md\n"
        "├── PLANO.md\n"
        "└── RELATORIO.md",
    )

    # =============== DECISOES DE PROJETO ===============
    add_heading(doc, "Decisões de projeto", 1)

    add_heading(
        doc,
        "Por que %rbp não pode usar leave no epílogo, apesar de o "
        "enunciado sugerir uma convenção clássica de frame pointer?",
        2,
    )
    add_body(
        doc,
        "Tentamos inicialmente usar `leave` (equivalente a `mov "
        "%rbp, %rsp; pop %rbp`) por parecer a forma mais idiomática/"
        "robusta de desfazer o prólogo, sem depender de recalcular L "
        "(número de locais) no ponto de saída. Os testes de "
        "equivalência semântica, porém, começaram a falhar com a "
        "pilha desbalanceada (`rsp` final diferente do inicial) em "
        "todo programa com pelo menos uma função com variável local "
        "— e só nesses casos.",
    )
    add_body(
        doc,
        "Rastreando manualmente, instrução por instrução, a execução "
        "simulada de `abs(17)` (que tem uma variável local, `y`), "
        "identificamos a causa: o prólogo adotado aqui faz `push "
        "%rbp` antes do `sub $8*L, %rsp`, e só copia `%rsp` para "
        "`%rbp` depois do `sub` — nessa ordem específica, `%rbp` "
        "termina apontando para o início do bloco de variáveis "
        "locais, não para o endereço onde o `%rbp` antigo foi salvo "
        "(que fica `8*L` bytes acima, mais perto do topo original da "
        "pilha). A instrução `leave` assume a convenção clássica, na "
        "qual `%rbp` aponta diretamente para esse slot salvo — o que "
        "só coincide com a convenção daqui quando `L == 0`.",
    )
    add_body(
        doc,
        "Usar `leave` estava, então, lendo de volta o valor da "
        "variável local como se fosse o `%rbp` salvo, corrompendo o "
        "frame pointer do chamador sempre que a função tinha locais. "
        "A correção foi reverter para as duas instruções literais do "
        "enunciado (`add $8*L, %rsp` seguido de `pop %rbp`), que não "
        "fazem essa suposição — cada uma desfaz exatamente o que o "
        "prólogo fez, na ordem inversa. Esse episódio é um bom "
        "lembrete de que uma “simplificação” aparentemente "
        "equivalente pode depender de uma invariante que o código, na "
        "verdade, não garante.",
    )

    add_heading(
        doc,
        "Por que os argumentos são empilhados em ordem inversa antes "
        "do call?",
        2,
    )
    add_body(
        doc,
        "Porque, combinado com o cálculo de deslocamento de "
        "parâmetro (`8*L + 16 + 8*i`, crescente com `i`), isso faz o "
        "primeiro parâmetro declarado ficar no deslocamento menor "
        "(mais perto do topo da pilha dentro do registro de ativação "
        "da função chamada) — a disposição exigida pelo exemplo da "
        "seção 6.1.3 do enunciado. Empilhar na ordem direta exigiria "
        "inverter também a fórmula de deslocamento.",
    )

    add_heading(
        doc,
        "Por que a análise semântica registra uma FunDecl na tabela "
        "de símbolos antes de verificar seu próprio corpo, e não "
        "depois?",
        2,
    )
    add_body(
        doc,
        "É exatamente essa ordem que permite recursão direta: se o "
        "nome da função só fosse inserido na tabela depois que seu "
        "corpo terminasse de ser verificado, uma chamada a si mesma "
        "dentro do próprio corpo seria erroneamente rejeitada como "
        "“função não declarada”. Como o processamento das "
        "declarações continua sendo sequencial (uma função só "
        "“existe” para as declarações seguintes, não para as "
        "anteriores), recursão mútua continua sendo rejeitada sem "
        "nenhum código adicional — é uma consequência direta da ordem "
        "escolhida, não um caso especial tratado à parte.",
    )

    add_heading(
        doc,
        "Por que o simulador de teste modela a pilha como um "
        "dicionário endereço → valor, com %rsp/%rbp como inteiros de "
        "verdade, em vez de continuar usando índices de lista como "
        "nas atividades anteriores?",
        2,
    )
    add_body(
        doc,
        "Porque, a partir desta atividade, o código gerado acessa "
        "memória por deslocamento relativo a `%rbp` (`16(%rbp)`, por "
        "exemplo) — um endereço calculado em tempo de execução, não "
        "um índice fixo conhecido estaticamente. Só um modelo com "
        "endereços “de verdade” (que crescem e decrescem exatamente "
        "como ponteiros reais) permite traduzir essas instruções "
        "quase literalmente, sem reescrever a semântica de "
        "deslocamento “por fora” do simulador.",
    )

    add_heading(
        doc,
        "Por que precisamos de uma pilha de retorno própria (dentro "
        "do mesmo dicionário de memória) para call/ret, em vez de "
        "usar a pilha de chamadas do próprio Python?",
        2,
    )
    add_body(
        doc,
        "Porque o programa Fun sendo simulado pode ser recursivo "
        "(fib), e o objetivo do simulador é validar que o código "
        "gerado gerencia a pilha corretamente por conta própria — "
        "inclusive o endereço de retorno, que no x86-64 real também "
        "mora na pilha do sistema, ao lado dos argumentos e "
        "variáveis locais. Delegar isso à pilha de chamadas do Python "
        "não testaria nada sobre o código gerado.",
    )

    # =============== DIFICULDADES ===============
    add_heading(doc, "Dificuldades", 1)
    add_body(
        doc,
        "Diferente das atividades anteriores, esta teve duas "
        "dificuldades genuínas — ambas descobertas (e resolvidas) "
        "pela própria suíte de testes, não por leitura do enunciado:",
    )
    add_bullet(
        doc,
        "O bug do leave, descrito em detalhe na seção “Decisões de "
        "projeto” acima: uma escolha de codegen que parecia "
        "estritamente mais robusta que a sugestão literal do "
        "enunciado, mas que na verdade dependia de uma invariante "
        "(`%rbp` apontando para o slot salvo) que a convenção de "
        "prólogo adotada aqui não garante quando a função tem "
        "variáveis locais. Só foi percebido porque o simulador de "
        "teste reproduz a semântica de `%rbp`/`%rsp` com fidelidade "
        "suficiente para expor a inconsistência — um caso claro de "
        "bug que só um teste de equivalência de comportamento (e não "
        "um teste de estrutura sintática do .s gerado) seria capaz de "
        "capturar.",
    )
    add_bullet(
        doc,
        "Um segundo bug, este no próprio simulador de teste: o “fim "
        "do bloco principal” era originalmente detectado comparando "
        "o program counter contra o índice onde termina a lista de "
        "instruções do main — mas esse índice é exatamente onde "
        "começa o rótulo da primeira função declarada, já que as "
        "funções são “penduradas” logo em seguida no .s gerado. Isso "
        "fazia o simulador confundir “um call acabou de saltar para "
        "dentro da primeira função” com “o main terminou por "
        "fallthrough natural”, interrompendo a simulação antes de a "
        "função sequer começar a executar — e deixando na pilha os "
        "argumentos e o endereço de retorno que nunca chegavam a ser "
        "desfeitos, o que se manifestava como o mesmo sintoma do bug "
        "anterior (pilha desbalanceada ao final). A correção foi "
        "parar de inferir o fim do bloco principal por posição no "
        "texto, e em vez disso tratar a própria linha "
        "`call imprime_num` (incluída no corpo simulado) como uma "
        "instrução sentinela de parada — o que só é verdade quando "
        "ela é de fato alcançada por execução sequencial, e nunca por "
        "um CALL (que sempre teria `imprime_num` como alvo inválido). "
        "Esse bug reforçou a importância de isolar, ao depurar uma "
        "falha de teste, se a causa está no código sendo testado ou "
        "no próprio arnês de teste — os dois produziam exatamente o "
        "mesmo sintoma (`rsp` final incorreto), mas por motivos "
        "completamente diferentes.",
    )

    doc.save(path)


if __name__ == "__main__":
    import os
    out = os.path.join(os.path.dirname(__file__), "RELATORIO.docx")
    build_document(out)
    print(f"gerado: {out}")
